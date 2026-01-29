 ################## ИМПОРТЫ ##################
import os, time, traceback, json, threading, re, httpx,  io
from typing import Optional, Dict, Any, List, Tuple
from zoneinfo import ZoneInfo
from fastapi import UploadFile, File, Form, FastAPI, APIRouter, HTTPException, Query, Body, Request, Depends
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response, PlainTextResponse
from starlette.exceptions import HTTPException as StarletteHTTPException
from datetime import date
from pydantic import BaseModel, EmailStr
from contextlib import asynccontextmanager
from openai import OpenAI
from datetime import datetime
from decimal import Decimal
from db import (
    jobs_create,
    jobs_get,
    jobs_list_by_decl,
    jobs_claim_next,
    jobs_finish_ok,
    jobs_finish_err,
    get_file,
    get_user_by_email,
    get_user_by_id,
    create_user,
    update_user,
    add_declaration,
    list_declarations,
    update_declaration,
    add_file,
    list_declaration_files,
    link_file_to_declaration,
    unlink_file_from_declaration,
    get_declaration_date,
    get_overrides,
    save_overrides,
    get_declaration_invoice_json,
    save_declaration_invoice_json,
    get_declaration_user_id,
    get_conn,
    list_active_tariff_plans,
    get_tariff_plan_by_code,
    payments_create_pending,
    payments_set_provider_payment_id,
    payments_get_by_id_and_user,
    payments_update_status_by_provider_id,
    credits_apply_purchase,
    credits_get_balance,
    credits_spend,          
    credits_get_status,
    credits_consume,
    tnved_requests_add,
    tnved_requests_list_by_user,
    payments_list_by_user,
)

from yandex_ocr import extract_text_with_meta
from openpyxl import load_workbook
from graph import extract_index 
import fitz
from lxml import etree
from docx import Document

import logging
from datetime import timedelta
from jose import jwt, JWTError
from passlib.context import CryptContext
from collections import defaultdict, deque

from xmlmap.ESADout_CU import (
    DocumentID as ESADout_CU_DocumentID,
    CustomsProcedure,
    CustomsModeCode,
    ElectronicDocumentSign,
    RecipientCountryCode,
    EECEDocHeaderAddInfo,
    ESADout_CU,
)

from xmlmap.ESADout_CUGoodsShipment import (
    ESADout_CUGoodsShipment,
    ESADout_CUConsignor,
    SubjectAddressDetails,
    ESADout_CUConsignee,
    ESADout_CUFinancialAdjustingResponsiblePerson,
    RFOrganizationFeatures,
    ESADout_CUDeclarant,
    ESADout_CUConsigment,
    BorderCustomsOffice,
    RUTransportMeans,
    ESADout_CUDepartureArrivalTransport,
    ESADout_CUBorderTransport,
    CUESADDeliveryTerms,
    ESADout_CUMainContractTerms,
    ESADout_CUGoodsLocation,
    RegisterDocumentIdDetails,
    GoodsGroupQuantity,
    GoodsGroupInformation,
    GoodsGroupDescription,
    Preferencii,
    DocumentPresentingDetails,
    ESADout_CUPresentedDocument,
    PackagePalleteInformation,
    ESADGoodsPackaging,
    ESADCustomsProcedure,
    SupplementaryGoodsQuantity,
    ESADout_CUGoods,
)




OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")  

YANDEX_API_KEY = os.getenv("YANDEX_API_KEY")  
YANDEX_FOLDER  = os.getenv("YANDEX_FOLDER_ID")

YANDEX_CLOUD_FOLDER = os.getenv("YANDEX_CLOUD_FOLDER")
YANDEX_CLOUD_MODEL = os.getenv("YANDEX_CLOUD_MODEL")

OFDATA_API_KEY = os.getenv("OFDATA_API_KEY")  
OFDATA_URL = "https://api.ofdata.ru/v2/company"



APP_ENV = os.getenv("APP_ENV", "dev").lower()  # dev|staging|prod
JWT_SECRET = os.getenv("JWT_SECRET", "")
JWT_ALG = os.getenv("JWT_ALG", "HS256")
ACCESS_TOKEN_EXPIRE_MIN = int(os.getenv("ACCESS_TOKEN_EXPIRE_MIN", "43200")) 


_rl = defaultdict(deque)
def rate_limit(key: str, limit: int, window_sec: int):
    now = time.time()
    q = _rl[key]
    while q and (now - q[0] > window_sec):
        q.popleft()
    if len(q) >= limit:
        raise HTTPException(429, "Too many requests")
    q.append(now)

def client_ip(request: Request) -> str:
    xff = request.headers.get("X-Forwarded-For")
    if xff:
        return xff.split(",")[0].strip()
    return request.client.host if request.client else "unknown"

def pdf_page_count(data: bytes) -> int:
    doc = fitz.open(stream=data, filetype="pdf")
    n = doc.page_count
    doc.close()
    return n

def gpt_client():
    return openai_client()

def openai_client() -> OpenAI:
    if not OPENAI_API_KEY:
        raise RuntimeError("OPENAI_API_KEY is not set")
    return OpenAI(api_key=OPENAI_API_KEY)

def build_prompt(doc_key: str, filename: str, extracted_text: str) -> str:
    return (
        f"Тип документа: {doc_key}.\n"
        "Задача: извлечь структуру данных из приведённого ниже текста и вывести результат "
        "СТРОГО в соответствии с JSON-схемой указанного типа документа. "
        "Заполняй только те поля, что присутствуют в схеме. "
        "Если значение отсутствует в тексте — укажи null. "
        "Не добавляй лишних ключей и не меняй формат.\n\n"

        "Общие правила обработки:\n"
        "1) Используй только информацию из источника. Никаких домыслов.\n"
        "2) Форматы: даты — ДД.ММ.ГГГГ; проценты — например, \"5%\"; числовые значения — без пробелов;\n"
        "   валюты — ISO-коды (CNY, USD, EUR и т.п.).\n"
        "3) Если в разделе оплаты (payment) встречаются выражения с корнем \"страхов*\", "
        "   относи такую услугу к категории страхования.\n"
        "4) В международных документах используй соответствия:\n"
        "     Seller = продавец = экспортер = отправитель товара\n"
        "     Buyer  = покупатель = импортер = получатель товара\n"
        "5) Если какое-либо значение указано как единое для группы товаров "
        "   (например, масса брутто 1 кг указана сразу на 4 позиции), "
        "   то равномерно распределяй указанное значение между всеми товарами.\n"
        "6) Строки вида \"ИТОГО\" или \"TOTAL\" не считай отдельными товарами — "
        "   такие строки не должны попадать в массив товаров.\n\n"

        "Требования к результату:\n"
        "— Выведи ответ строго между тегами <#START_JSON#> и <#END_JSON#>.\n"
        "— Используй только синтаксически корректный JSON.\n"
        "— Соблюдай точное соответствие JSON-схеме данного документа.\n\n"

        "<source>\n"
        f"{extracted_text[:20000]}\n"
        "</source>"
    )


################## JSON Schemas ##################
def json_schema_for(doc_key: str) -> dict:
    if doc_key == "invoice":
        return {
            "name": "InvoiceV1",
            "schema": {
                "type": "object",
                "additionalProperties": False,
                "properties": {
                    "Общая информация": {
                        "type": "object",
                        "additionalProperties": False,
                        "properties": {
                            "Номер инвойса": {"type": ["string","null"]},
                            "Дата инвойса": {"type": ["string","null"]},
                            "Срок оплаты":  {"type": ["string","null"]},
                            "Условия поставки (Incoterms)": {"type": ["string","null"]}
                        },
                        "required": ["Номер инвойса","Дата инвойса","Срок оплаты","Условия поставки (Incoterms)"]
                    },
                    "Отправитель": {
                        "type": "object",
                        "additionalProperties": False,
                        "properties": {
                            "Название компании": {"type": ["string","null"]},
                            "Юридический адрес": {
                                "type": "object",
                                "additionalProperties": False,
                                "properties": {
                                    "Полностью": {"type": ["string","null"]},
                                    "Страна": {"type": ["string","null"]},
                                    "Регион/Область": {"type": ["string","null"]},
                                    "Город": {"type": ["string","null"]},
                                    "Улица": {"type": ["string","null"]},
                                    "Номер дома": {"type": ["string","null"]}
                                },
                                "required": ["Полностью","Страна","Регион/Область","Город","Улица","Номер дома"]
                            },
                            "Страна": {"type": ["string","null"]},
                            "ИНН": {"type": ["string","null"]},
                            "КПП": {"type": ["string","null"]},
                            "Контакты": {
                                "type": "object",
                                "additionalProperties": False,
                                "properties": {
                                    "Контактное лицо": {"type": ["string","null"]},
                                    "Телефон": {"type": ["string","null"]},
                                    "Почта": {"type": ["string","null"]}
                                },
                                "required": ["Контактное лицо","Телефон","Почта"]
                            }
                        },
                        "required": ["Название компании","Юридический адрес","Страна","ИНН","КПП","Контакты"]
                    },
                    "Получатель": { "$ref": "#/properties/Отправитель" },
                    "Товары": {
                        "type": "array",
                        "items": {
                            "type": "object",
                            "additionalProperties": False,
                            "properties": {
                                "Наименование": {"type": ["string","null"]},
                                "Количество": {"type": ["string","null","number"]},
                                "Единица измерения": {"type": ["string","null"]},
                                "Цена": {"type": ["string","null","number"]},
                                "Валюта": {"type": ["string","null"]},
                                "Стоимость": {"type": ["string","null","number"]},
                                "Страна-производитель": {"type": ["string","null"]},
                                "Код ТНВЭД": {"type": ["string","null"]},
                                "Пошлина": {"type": ["string","null"]},
                                "НДС": {"type": ["string","null"]},
                                "Дополнительная информация": {"type": ["string","null"]},
                                "Техническое описание": {"type": ["string","null"]}
                            },
                            "required": ["Наименование","Количество","Единица измерения","Цена","Валюта","Стоимость",
                                         "Страна-производитель","Код ТНВЭД","Пошлина","НДС",
                                         "Дополнительная информация","Техническое описание"]
                        }
                    }
                },
                "required": ["Общая информация","Отправитель","Получатель","Товары"]
            }
        }

    if doc_key == "contract":
        return {
            "name": "ContractV1",
            "schema": {
                "type": "object",
                "additionalProperties": False,
                "$defs": {
                    "Address": {
                        "type": "object",
                        "additionalProperties": False,
                        "properties": {
                            "Полностью": {"type": ["string","null"]},
                            "Страна": {"type": ["string","null"]},
                            "Регион/Область": {"type": ["string","null"]},
                            "Город": {"type": ["string","null"]},
                            "Улица": {"type": ["string","null"]},
                            "Номер дома": {"type": ["string","null"]}
                        },
                        "required": ["Полностью","Страна","Регион/Область","Город","Улица","Номер дома"]
                    },
                    "Contacts": {
                        "type": "object",
                        "additionalProperties": False,
                        "properties": {
                            "Контактное лицо": {"type": ["string","null"]},
                            "Телефон": {"type": ["string","null"]},
                            "Почта": {"type": ["string","null"]}
                        },
                        "required": ["Контактное лицо","Телефон","Почта"]
                    },
                    "Bank": {
                        "type": "object",
                        "additionalProperties": False,
                        "properties": {
                            "Наименование банка": {"type": ["string","null"]},
                            "Корреспондентский счет": {"type": ["string","null"]},
                            "Номер счета": {"type": ["string","null"]}
                        },
                        "required": ["Наименование банка","Корреспондентский счет","Номер счета"]
                    },
                    "PartyWithBank": {
                        "type": "object",
                        "additionalProperties": False,
                        "properties": {
                            "Название компании": {"type": ["string","null"]},
                            "Юридический адрес": {"$ref": "#/$defs/Address"},
                            "Страна": {"type": ["string","null"]},
                            "ИНН": {"type": ["string","null"]},
                            "КПП": {"type": ["string","null"]},
                            "ОГРН": {"type": ["string","null"]},
                            "Банк": {"$ref": "#/$defs/Bank"},
                            "Контакты": {"$ref": "#/$defs/Contacts"}
                        },
                        "required": ["Название компании","Юридический адрес","Страна","ИНН","КПП","ОГРН","Банк","Контакты"]
                    }
                },
                "properties": {
                    "Общая информация": {
                        "type": "object",
                        "additionalProperties": False,
                        "properties": {
                            "Номер контракта": {"type": ["string","null"]},
                            "Дата заключения": {"type": ["string","null"]},
                            "Дата прекращения": {"type": ["string","null"]},
                            "Стороны": {
                                "type": "object",
                                "additionalProperties": False,
                                "properties": {
                                    "Отправитель": {"$ref": "#/$defs/PartyWithBank"},
                                    "Получатель": {"$ref": "#/$defs/PartyWithBank"}
                                },
                                "required": ["Отправитель","Получатель"]
                            }
                        },
                        "required": ["Номер контракта","Дата заключения","Дата прекращения","Стороны"]
                    },
                    "Декларант": {"$ref": "#/$defs/PartyWithBank"},
                    "Оплата контракта": {
                        "type": "object",
                        "additionalProperties": False,
                        "properties": {
                            "Общая сумма": {"type": ["string","number","null"]},
                            "Валюта": {"type": ["string","null"]},
                            "Предоплата": {
                                "type": "object",
                                "additionalProperties": False,
                                "properties": {
                                    "Тип предоплаты": {"type": ["string","null"]},
                                    "Размер предоплаты": {"type": ["string","number","null"]},
                                    "Сроки предоплаты": {"type": ["string","null"]}
                                },
                                "required": ["Тип предоплаты","Размер предоплаты","Сроки предоплаты"]
                            }
                        },
                        "required": ["Общая сумма","Валюта","Предоплата"]
                    },
                    "Поставка": {
                        "type": "object",
                        "additionalProperties": False,
                        "properties": {
                            "Условия поставки (Incoterms)": {"type": ["string","null"]},
                            "Место поставки": {"type": ["string","null"]},
                            "Срок поставки": {"type": ["string","null"]}
                        },
                        "required": ["Условия поставки (Incoterms)","Место поставки","Срок поставки"]
                    }
                },
                "required": ["Общая информация","Декларант","Оплата контракта","Поставка"]
            }
        }

    if doc_key == "packing":
        return {
            "name": "PackingV1",
            "schema": {
                "type": "object",
                "additionalProperties": False,
                "$defs": {
                    "Address": {
                        "type": "object",
                        "additionalProperties": False,
                        "properties": {
                            "Полностью": {"type": ["string","null"]},
                            "Страна": {"type": ["string","null"]},
                            "Регион/Область": {"type": ["string","null"]},
                            "Город": {"type": ["string","null"]},
                            "Улица": {"type": ["string","null"]},
                            "Номер дома": {"type": ["string","null"]}
                        },
                        "required": ["Полностью","Страна","Регион/Область","Город","Улица","Номер дома"]
                    },
                    "Contacts": {
                        "type": "object",
                        "additionalProperties": False,
                        "properties": {
                            "Контактное лицо": {"type": ["string","null"]},
                            "Телефон": {"type": ["string","null"]},
                            "Почта": {"type": ["string","null"]}
                        },
                        "required": ["Контактное лицо","Телефон","Почта"]
                    },
                    "PartyWithIds": {
                        "type": "object",
                        "additionalProperties": False,
                        "properties": {
                            "Название компании": {"type": ["string","null"]},
                            "Юридический адрес": {"$ref": "#/$defs/Address"},
                            "Страна": {"type": ["string","null"]},
                            "ИНН": {"type": ["string","null"]},
                            "КПП": {"type": ["string","null"]},
                            "Контакты": {"$ref": "#/$defs/Contacts"}
                        },
                        "required": ["Название компании","Юридический адрес","Страна","ИНН","КПП","Контакты"]
                    },
                    "Place": {
                        "type": "object",
                        "additionalProperties": False,
                        "properties": {
                            "Адрес": {"type": ["string","null"]},
                            "Страна": {"type": ["string","null"]},
                            "Дата": {"type": ["string","null"]}
                        },
                        "required": ["Адрес","Страна","Дата"]
                    },
                    "Item": {
                        "type": "object",
                        "additionalProperties": False,
                        "properties": {
                            "Наименование": {"type": ["string","null"]},
                            "Количество": {"type": ["string","number","null"]},
                            "Единица измерения": {"type": ["string","null"]},
                            "Цена": {"type": ["string","number","null"]},
                            "Валюта": {"type": ["string","null"]},
                            "Стоимость": {"type": ["string","number","null"]},
                            "Страна-производитель": {"type": ["string","null"]},
                            "Код ТНВЭД": {"type": ["string","null"]},
                            "Пошлина": {"type": ["string","null"]},
                            "НДС": {"type": ["string","null"]},
                            "Габариты": {"type": ["string","null"]},
                            "Маркировка": {"type": ["string","null"]},
                            "Масса брутто": {"type": ["string","number","null"]},
                            "Масса нетто": {"type": ["string","number","null"]},
                            "Количество мест": {"type": ["string","number","null"]}
                        },
                        "required": ["Наименование","Количество","Единица измерения","Цена","Валюта","Стоимость",
                                     "Страна-производитель","Код ТНВЭД","Пошлина","НДС",
                                     "Габариты","Маркировка","Масса брутто","Масса нетто","Количество мест"]
                    }
                },
                "properties": {
                    "Общая информация": {
                        "type": "object",
                        "additionalProperties": False,
                        "properties": {
                            "Номер упаковочного листа (Packing List)": {"type": ["string","null"]},
                            "Дата упаковочного листа (Packing List)": {"type": ["string","null"]},
                            "Номер контракта": {"type": ["string","null"]},
                            "Номер инвойса(счета)": {"type": ["string","null"]}
                        },
                        "required": ["Номер упаковочного листа (Packing List)","Дата упаковочного листа (Packing List)",
                                     "Номер контракта","Номер инвойса(счета)"]
                    },
                    "Перевозка": {
                        "type": "object",
                        "additionalProperties": False,
                        "properties": {
                            "Отправитель": {"$ref": "#/$defs/PartyWithIds"},
                            "Получатель": {"$ref": "#/$defs/PartyWithIds"},
                            "Место погрузки": {"$ref": "#/$defs/Place"},
                            "Место разгрузки": {"$ref": "#/$defs/Place"}
                        },
                        "required": ["Отправитель","Получатель","Место погрузки","Место разгрузки"]
                    },
                    "Товары": {
                        "type": "array",
                        "items": {"$ref": "#/$defs/Item"}
                    }
                },
                "required": ["Общая информация","Перевозка","Товары"]
            }
        }

    if doc_key == "payment":
        return {
        "name": "BillLogisticsV1",
        "schema": {
            "type": "object",
            "additionalProperties": False,
            "$defs": {
            "Address": {
                "type": "object",
                "additionalProperties": False,
                "properties": {
                "Полностью": { "type": ["string","null"] },
                "Страна": { "type": ["string","null"] },
                "Регион/Область": { "type": ["string","null"] },
                "Город": { "type": ["string","null"] },
                "Улица": { "type": ["string","null"] },
                "Номер дома": { "type": ["string","null"] }
                },
                "required": ["Полностью","Страна","Регион/Область","Город","Улица","Номер дома"]
            },
            "Party": {
                "type": "object",
                "additionalProperties": False,
                "properties": {
                "Название компании": { "type": ["string","null"] },
                "Юридический адрес": { "$ref": "#/$defs/Address" },
                "Страна": { "type": ["string","null"] },
                "ИНН": { "type": ["string","null"] },
                "КПП": { "type": ["string","null"] },
                "Контакты": {
                    "type": "object",
                    "additionalProperties": False,
                    "properties": {
                    "Контактное лицо": { "type": ["string","null"] },
                    "Телефон": { "type": ["string","null"] },
                    "Почта": { "type": ["string","null"] }
                    },
                    "required": ["Контактное лицо","Телефон","Почта"]
                }
                },
                "required": ["Название компании","Юридический адрес","Страна","ИНН","КПП","Контакты"]
            },
            "MoneyLine": {
                "type": "object",
                "additionalProperties": False,
                "properties": {
                "Описание": { "type": ["string","null"] },
                "Кол-во": { "type": ["number","string","null"] },
                "Цена": { "type": ["number","string","null"] },
                "Сумма": { "type": ["number","string","null"] },
                "Номер инвойса": { "type": ["string","null"] }
                },
                "required": ["Описание","Кол-во","Цена","Сумма","Номер инвойса"]
            }
            },
            "properties": {
            "Общая информация": {
                "type": "object",
                "additionalProperties": False,
                "properties": {
                "Номер счета": { "type": ["string","null"] },
                "Дата счета": { "type": ["string","null"] },
                "Основание": { "type": ["string","null"] },   
                "Валюта документа": { "type": ["string","null"] },
                "Язык": { "type": ["string","null"] },
                "Цена включает НДС": { "type": ["boolean","null"] }
                },
                "required": ["Номер счета","Дата счета","Основание","Валюта документа","Язык","Цена включает НДС"]
            },
            "Банк получателя": {
                "type": "object",
                "additionalProperties": False,
                "properties": {
                "Банк": { "type": ["string","null"] },
                "БИК": { "type": ["string","null"] },
                "Корр.счёт": { "type": ["string","null"] },
                "Расч.счёт": { "type": ["string","null"] }
                },
                "required": ["Банк","БИК","Корр.счёт","Расч.счёт"]
            },
            "Поставщик (Исполнитель)": { "$ref": "#/$defs/Party" },
            "Покупатель (Заказчик)": { "$ref": "#/$defs/Party" },

            "Перевозка": {
                "type": "array",
                "items": {
                "type": "object",
                "additionalProperties": False,
                "properties": {
                    "Маршрут": {
                    "type": "object",
                    "additionalProperties": False,
                    "properties": {
                        "Откуда": { "type": ["string","null"] },
                        "Куда": { "type": ["string","null"] }
                    },
                    "required": ["Откуда","Куда"]
                    },     
                    "Услуга": { "$ref": "#/$defs/MoneyLine" }
                },
                "required": ["Маршрут","Услуга"]
                }
            },

            "Страхование": {
                "type": "array",
                "items": {
                "type": "object",
                "additionalProperties": False,
                "properties": {
                                        "Маршрут": {
                    "type": "object",
                    "additionalProperties": False,
                    "properties": {
                        "Откуда": { "type": ["string","null"] },
                        "Куда": { "type": ["string","null"] }
                    },
                    "required": ["Откуда","Куда"]
                    },
                    "Услуга": { "$ref": "#/$defs/MoneyLine" }
                },
                "required": ["Маршрут","Услуга"]
                }
            }
            }
        }
        }

    if doc_key == "transport_road":
        return {
            "name": "TransportV1",
            "schema": {
                "type": "object",
                "additionalProperties": False,
                "$defs": {
                    "Address": {
                        "type": "object",
                        "additionalProperties": False,
                        "properties": {
                            "Полностью": {"type": ["string","null"]},
                            "Страна": {"type": ["string","null"]},
                            "Регион/Область": {"type": ["string","null"]},
                            "Город": {"type": ["string","null"]},
                            "Улица": {"type": ["string","null"]},
                            "Номер дома": {"type": ["string","null"]}
                        },
                        "required": ["Полностью","Страна","Регион/Область","Город","Улица","Номер дома"]
                    },
                    "Contacts": {
                        "type": "object",
                        "additionalProperties": False,
                        "properties": {
                            "Контактное лицо": {"type": ["string","null"]},
                            "Телефон": {"type": ["string","null"]},
                            "Почта": {"type": ["string","null"]}
                        },
                        "required": ["Контактное лицо","Телефон","Почта"]
                    },
                    "PartyWithIds": {
                        "type": "object",
                        "additionalProperties": False,
                        "properties": {
                            "Название компании": {"type": ["string","null"]},
                            "Юридический адрес": {"$ref": "#/$defs/Address"},
                            "Страна": {"type": ["string","null"]},
                            "ИНН": {"type": ["string","null"]},
                            "КПП": {"type": ["string","null"]},
                            "Контакты": {"$ref": "#/$defs/Contacts"}
                        },
                        "required": ["Название компании","Юридический адрес","Страна","ИНН","КПП","Контакты"]
                    },
                    "Place": {
                        "type": "object",
                        "additionalProperties": False,
                        "properties": {
                            "Адрес": {"type": ["string","null"]},
                            "Страна": {"type": ["string","null"]},
                            "Дата": {"type": ["string","null"]}
                        },
                        "required": ["Адрес","Страна","Дата"]
                    },
                    "Item": {
                        "type": "object",
                        "additionalProperties": False,
                        "properties": {
                            "Наименование": {"type": ["string","null"]},
                            "Количество": {"type": ["string","number","null"]},
                            "Единица измерения": {"type": ["string","null"]},
                            "Цена": {"type": ["string","number","null"]},
                            "Валюта": {"type": ["string","null"]},
                            "Стоимость": {"type": ["string","number","null"]},
                            "Страна-производитель": {"type": ["string","null"]},
                            "Код ТНВЭД": {"type": ["string","null"]},
                            "Пошлина": {"type": ["string","null"]},
                            "НДС": {"type": ["string","null"]},
                            "Габариты": {"type": ["string","null"]},
                            "Маркировка": {"type": ["string","null"]},
                            "Масса брутто": {"type": ["string","number","null"]},
                            "Масса нетто": {"type": ["string","number","null"]},
                            "Количество мест": {"type": ["string","number","null"]}
                        },
                        "required": ["Наименование","Количество","Единица измерения","Цена","Валюта","Стоимость",
                                     "Страна-производитель","Код ТНВЭД","Пошлина","НДС",
                                     "Габариты","Маркировка","Масса брутто","Масса нетто","Количество мест"]
                    }
                },
                "properties": {
                    "Общая информация": {
                        "type": "object",
                        "additionalProperties": False,
                        "properties": {
                            "Номер накладной": {"type": ["string","null"]},
                            "Дата накладной": {"type": ["string","null"]}
                        },
                        "required": ["Номер накладной","Дата накладной"]
                    },
                    "Перевозка": {
                        "type": "object",
                        "additionalProperties": False,
                        "properties": {
                            "Отправитель": {"$ref": "#/$defs/PartyWithIds"},
                            "Получатель": {"$ref": "#/$defs/PartyWithIds"},
                            "Место погрузки": {"$ref": "#/$defs/Place"},
                            "Место разгрузки": {"$ref": "#/$defs/Place"},
                            "Условия": {
                                "type": "object",
                                "additionalProperties": False,
                                "properties": {
                                    "Условия поставки (Incoterms)": {"type": ["string","null"]}
                                },
                                "required": ["Условия поставки (Incoterms)"]
                            },
                            "Регистрационный номер": {
                                "type": "object",
                                "additionalProperties": False,
                                "properties": {
                                    "Тягач": {"type": ["string","array","null"]},
                                    "Прицеп": {"type": ["string","array","null"]}
                                },
                                "required": ["Тягач","Прицеп"]
                            }
                        },
                        "required": ["Отправитель","Получатель","Место погрузки","Место разгрузки","Условия","Регистрационный номер"]
                    },
                    "Таможенный пост": {
                        "type": "object",
                        "additionalProperties": False,
                        "properties": {
                            "Код ТП": {"type": ["string","null"]},
                            "Наименование таможенного поста": {"type": ["string","null"]},
                            "Адрес": {"type": ["string","null"]},
                            "Номер лицензии (License/Lic.)": {"type": ["string","null"]}
                        },
                        "required": ["Код ТП","Наименование таможенного поста","Адрес","Номер лицензии (License/Lic.)"]
                    },
                    "Товары": {
                        "type": "array",
                        "items": {"$ref": "#/$defs/Item"}
                    }
                },
                "required": ["Общая информация","Перевозка","Таможенный пост","Товары"]
            }
        }
    
    if doc_key == "transport_air":
        return {
            "name": "AirWaybillV1",
            "schema": {
                "type": "object",
                "additionalProperties": False,
                "$defs": {
                "Address": {
                    "type": "object",
                    "additionalProperties": False,
                    "properties": {
                        "Полностью": { "type": ["string","null"] },
                        "Страна": { "type": ["string","null"] },
                        "Регион/Область": { "type": ["string","null"] },
                        "Город": { "type": ["string","null"] },
                        "Улица": { "type": ["string","null"] },
                        "Номер дома": { "type": ["string","null"] }
                        },
                    "required": ["Полностью","Страна","Регион/Область","Город","Улица","Номер дома"]
                },
                "Contacts": {
                    "type": "object",
                    "additionalProperties": False,
                    "properties": {
                    "Контактное лицо": { "type": ["string","null"] },
                    "Телефон": { "type": ["string","null"] },
                    "Почта": { "type": ["string","null"] }
                    },
                    "required": ["Контактное лицо","Телефон","Почта"]
                },
                "PartyWithIds": {
                    "type": "object",
                    "additionalProperties": False,
                    "properties": {
                    "Название компании": { "type": ["string","null"] },
                    "Юридический адрес": { "$ref": "#/$defs/Address" },
                    "Страна": { "type": ["string","null"] },
                    "ИНН": { "type": ["string","null"] },
                    "КПП": { "type": ["string","null"] },
                    "Контакты": { "$ref": "#/$defs/Contacts" }
                    },
                    "required": ["Название компании","Юридический адрес","Страна","ИНН","КПП","Контакты"]
                },
                "AirportPlace": {
                    "type": "object",
                    "additionalProperties": False,
                    "properties": {
                    "Аэропорт": { "type": ["string","null"] },
                    "Город": { "type": ["string","null"] },
                    "Страна": { "type": ["string","null"] },
                    "Дата/время": { "type": ["string","null"] }
                    },
                    "required": ["Аэропорт","Город","Страна","Дата/время"]
                },
                "Item": {
                    "type": "object",
                    "additionalProperties": False,
                    "properties": {
                    "Наименование": { "type": ["string","null"] },
                    "Количество": { "type": ["string","number","null"] },
                    "Единица измерения": { "type": ["string","null"] },
                    "Цена": { "type": ["string","number","null"] },
                    "Валюта": { "type": ["string","null"] },
                    "Стоимость": { "type": ["string","number","null"] },
                    "Страна-производитель": { "type": ["string","null"] },
                    "Код ТНВЭД": { "type": ["string","null"] },
                    "Пошлина": { "type": ["string","null"] },
                    "НДС": { "type": ["string","null"] },
                    "Габариты": { "type": ["string","null"] },
                    "Масса брутто": { "type": ["string","number","null"] },
                    "Масса нетто": { "type": ["string","number","null"] },
                    "Количество мест": { "type": ["string","number","null"] },
                    },
                    "required": [
                    "Наименование","Количество","Единица измерения","Цена","Валюта","Стоимость",
                    "Страна-производитель","Код ТНВЭД","Пошлина","НДС",
                    "Габариты","Масса брутто","Масса нетто","Количество мест"
                    ]
                }
                },
                "properties": {
                "Общая информация": {
                    "type": "object",
                    "additionalProperties": False,
                    "properties": {
                    "Номер AWB": { "type": ["string","null"] },
                    "Дата AWB": { "type": ["string","null"] }
                    },
                    "required": ["Номер AWB","Дата AWB"]
                },
                "Перевозка": {
                    "type": "object",
                    "additionalProperties": False,
                    "properties": {
                    "Отправитель": { "$ref": "#/$defs/PartyWithIds" },
                    "Получатель": { "$ref": "#/$defs/PartyWithIds" },
                    "Аэропорт отправления": { "$ref": "#/$defs/AirportPlace" },
                    "Аэропорт назначения": { "$ref": "#/$defs/AirportPlace" },
                    "Перевозчик": {
                        "type": "object",
                        "additionalProperties": False,
                        "properties": {
                        "Авиакомпания": { "type": ["string","null"] },
                        "Номер рейса": { "type": ["string","null"] },
                        "Маршрут": { "type": ["string","null"] }
                        },
                        "required": ["Авиакомпания","Номер рейса","Маршрут"]
                    },
                    "Условия": {
                        "type": "object",
                        "additionalProperties": False,
                        "properties": {
                        "Условия поставки (Incoterms)": { "type": ["string","null"] }
                        },
                        "required": ["Условия поставки (Incoterms)"]
                    },
                    "Оплата фрахта": {
                        "type": "object",
                        "additionalProperties": False,
                        "properties": {
                        "Тип оплаты": { "type": ["string","null"] },
                        "Валюта": { "type": ["string","null"] }
                        },
                        "required": ["Тип оплаты","Валюта"]
                    }
                    },
                    "required": [
                    "Отправитель","Получатель",
                    "Аэропорт отправления","Аэропорт назначения",
                    "Перевозчик","Оплата фрахта","Условия"
                    ]
                },
                "Таможенный пост": {
                    "type": "object",
                    "additionalProperties": False,
                    "properties": {
                    "Код ТП": { "type": ["string","null"] },
                    "Наименование таможенного поста": { "type": ["string","null"] },
                    "Адрес": { "type": ["string","null"] },
                    "Номер лицензии (License/Lic.)": { "type": ["string","null"] }
                    },
                    "required": ["Код ТП","Наименование таможенного поста","Адрес","Номер лицензии (License/Lic.)"]
                },
                "Товары": {
                    "type": "array",
                    "items": { "$ref": "#/$defs/Item" }
                }
                },
                "required": ["Общая информация","Перевозка","Таможенный пост","Товары"]
            }
        }

    if doc_key == "transport_sea":
        return {
            "name": "BillOfLadingV1",
            "schema": {
                "type": "object",
                "additionalProperties": False,
                "$defs": {
                    "Address": {
                        "type": "object",
                        "additionalProperties": False,
                        "properties": {
                            "Полностью": {"type": ["string", "null"]},
                            "Страна": {"type": ["string", "null"]},
                            "Регион/Область": {"type": ["string", "null"]},
                            "Город": {"type": ["string", "null"]},
                            "Улица": {"type": ["string", "null"]},
                            "Номер дома": {"type": ["string", "null"]}
                        },
                        "required": ["Полностью","Страна","Регион/Область","Город","Улица","Номер дома"]
                    },
                    "Contacts": {
                        "type": "object",
                        "additionalProperties": False,
                        "properties": {
                            "Контактное лицо": {"type": ["string", "null"]},
                            "Телефон": {"type": ["string", "null"]},
                            "Почта": {"type": ["string", "null"]}
                        },
                        "required": ["Контактное лицо", "Телефон", "Почта"]
                    },
                    "PartyWithIds": {
                        "type": "object",
                        "additionalProperties": False,
                        "properties": {
                            "Название компании": {"type": ["string", "null"]},
                            "Юридический адрес": {"$ref": "#/$defs/Address"},
                            "Страна": {"type": ["string", "null"]},
                            "ИНН": {"type": ["string", "null"]},
                            "КПП": {"type": ["string", "null"]},
                            "Контакты": {"$ref": "#/$defs/Contacts"}
                        },
                        "required": ["Название компании","Юридический адрес","Страна",
                            "ИНН","КПП","Контакты"]
                    },
                    "PortPlace": {
                        "type": "object",
                        "additionalProperties": False,
                        "properties": {
                            "Порт (UN/LOCODE)": {"type": ["string", "null"]},
                            "Город": {"type": ["string", "null"]},
                            "Страна": {"type": ["string", "null"]},
                            "Дата/время": {"type": ["string", "null"]}
                        },
                        "required": ["Порт (UN/LOCODE)","Город","Страна","Дата/время"]
                    },
                    "Item": {
                        "type": "object",
                        "additionalProperties": False,
                        "properties": {
                            "Наименование": {"type": ["string", "null"]},
                            "Количество": {"type": ["string", "number", "null"]},
                            "Единица измерения": {"type": ["string", "null"]},
                            "Цена": {"type": ["string", "number", "null"]},
                            "Валюта": {"type": ["string", "null"]},
                            "Стоимость": {"type": ["string", "number", "null"]},
                            "Страна-производитель": {"type": ["string", "null"]},
                            "Код ТНВЭД": {"type": ["string", "null"]},
                            "Пошлина": {"type": ["string", "null"]},
                            "НДС": {"type": ["string", "null"]},
                            "Габариты": {"type": ["string", "null"]},
                            "Масса брутто": {"type": ["string", "number", "null"]},
                            "Масса нетто": {"type": ["string", "number", "null"]},
                            "Количество мест": {"type": ["string", "number", "null"]},
                            "Номер контейнера": {"type": ["string", "null"]},
                            "Тип контейнера": {"type": ["string", "null"]},
                            "Пломбы": {"type": ["string", "array", "null"]}
                        },
                        "required": ["Наименование","Количество","Единица измерения","Цена",
                            "Валюта","Стоимость","Страна-производитель","Код ТНВЭД",
                            "Пошлина","НДС","Габариты","Масса брутто","Масса нетто",
                            "Количество мест","Номер контейнера","Тип контейнера","Пломбы"]
                    }
                },
                "properties": {
                    "Общая информация": {
                        "type": "object",
                        "additionalProperties": False,
                        "properties": {
                            "Номер B/L": {"type": ["string", "null"]},
                            "Вид B/L": {"type": ["string", "null"]},
                            "Дата B/L": {"type": ["string", "null"]}
                        },
                        "required": ["Номер B/L", "Вид B/L", "Дата B/L"]
                    },
                    "Перевозка": {
                        "type": "object",
                        "additionalProperties": False,
                        "properties": {
                            "Отправитель": {"$ref": "#/$defs/PartyWithIds"},
                            "Получатель": {"$ref": "#/$defs/PartyWithIds"},
                            "Перевозчик": {"type": ["string", "null"]},
                            "Судно": {
                                "type": "object",
                                "additionalProperties": False,
                                "properties": {
                                    "Название судна": {"type": ["string", "null"]},
                                    "Рейс (Voyage)": {"type": ["string", "null"]}
                                },
                                "required": ["Название судна", "Рейс (Voyage)"]
                            },
                            "Условия": {
                                "type": "object",
                                "additionalProperties": False,
                                "properties": {
                                    "Условия поставки (Incoterms)": {"type": ["string", "null"]}
                                },
                                "required": ["Условия поставки (Incoterms)"]
                            },
                        },
                        "required": [
                            "Отправитель",
                            "Получатель",
                            "Судно",
                            "Условия"]
                    },
                    "Таможенный пост": {
                        "type": "object",
                        "additionalProperties": False,
                        "properties": {
                            "Код ТП": {"type": ["string", "null"]},
                            "Наименование таможенного поста": {"type": ["string", "null"]},
                            "Адрес": {"type": ["string", "null"]},
                            "Номер лицензии (License/Lic.)": {"type": ["string", "null"]}
                        },
                        "required": [
                            "Код ТП",
                            "Наименование таможенного поста",
                            "Адрес",
                            "Номер лицензии (License/Lic.)"
                        ]
                    },
                    "Товары": {
                        "type": "array",
                        "items": {"$ref": "#/$defs/Item"}
                    }
                },
                "required": [
                    "Общая информация",
                    "Перевозка",
                    "Таможенный пост",
                    "Товары"
                ]
            }
        }

    if doc_key == "transport_rail":
        return {
        "name": "RailConsignmentV1",
        "schema": {
            "type": "object",
            "additionalProperties": False,
            "$defs": {
            "Address": {
                "type": "object",
                "additionalProperties": False,
                "properties": {
                "Полностью": {"type": ["string","null"]},
                "Страна": {"type": ["string","null"]},
                "Регион/Область": {"type": ["string","null"]},
                "Город": {"type": ["string","null"]},
                "Улица": {"type": ["string","null"]},
                "Номер дома": {"type": ["string","null"]}
                },
                "required": ["Полностью","Страна","Регион/Область","Город","Улица","Номер дома"]
            },
            "Contacts": {
                "type": "object",
                "additionalProperties": False,
                "properties": {
                "Контактное лицо": {"type": ["string","null"]},
                "Телефон": {"type": ["string","null"]},
                "Почта": {"type": ["string","null"]}
                },
                "required": ["Контактное лицо","Телефон","Почта"]
            },
            "PartyWithIds": {
                "type": "object",
                "additionalProperties": False,
                "properties": {
                "Название компании": {"type": ["string","null"]},
                "Юридический адрес": {"$ref": "#/$defs/Address"},
                "Страна": {"type": ["string","null"]},
                "ИНН": {"type": ["string","null"]},
                "КПП": {"type": ["string","null"]},
                "Контакты": {"$ref": "#/$defs/Contacts"}
                },
                "required": ["Название компании","Юридический адрес","Страна","ИНН","КПП","Контакты"]
            },
            "StationPlace": {
                "type": "object",
                "additionalProperties": False,
                "properties": {
                "Станция": {"type": ["string","null"]},
                "Код станции": {"type": ["string","null"]},
                "Страна": {"type": ["string","null"]},
                "Дата/время": {"type": ["string","null"]}
                },
                "required": ["Станция","Код станции","Страна","Дата/время"]
            },
            # "RailUnit": {
            #     "type": "object",
            #     "additionalProperties": False,
            #     "properties": {
            #     "Тип единицы": {"type": ["string","null"]},
            #     "Номер": {"type": ["string","null"]},
            #     "Грузоподъёмность": {"type": ["string","number","null"]},
            #     "Масса тары": {"type": ["string","number","null"]},
            #     "Масса брутто": {"type": ["string","number","null"]}
            #     },
            #     "required": ["Тип единицы","Номер","Масса брутто"]
            # },
            "Item": {
                "type": "object",
                "additionalProperties": False,
                "properties": {
                "Наименование": {"type": ["string","null"]},
                "Количество": {"type": ["string","number","null"]},
                "Единица измерения": {"type": ["string","null"]},
                "Цена": {"type": ["string","number","null"]},
                "Валюта": {"type": ["string","null"]},
                "Стоимость": {"type": ["string","number","null"]},
                "Страна-производитель": {"type": ["string","null"]},
                "Код ТНВЭД": {"type": ["string","null"]},
                "Пошлина": {"type": ["string","null"]},
                "НДС": {"type": ["string","null"]},
                "Габариты": {"type": ["string","null"]},
                "Масса брутто": {"type": ["string","number","null"]},
                "Масса нетто": {"type": ["string","number","null"]},
                "Количество мест": {"type": ["string","number","null"]}
                },
                "required": [
                "Наименование","Количество","Единица измерения","Цена","Валюта","Стоимость",
                "Страна-производитель","Код ТНВЭД","Пошлина","НДС",
                "Габариты","Масса брутто","Масса нетто","Количество мест"
                ]
            }
            },
            "properties": {
            "Общая информация": {
                "type": "object",
                "additionalProperties": False,
                "properties": {
                "Номер ЖД накладной": {"type": ["string","null"]},
                "Дата накладной": {"type": ["string","null"]}
                },
                "required": ["Номер ЖД накладной","Дата накладной"]
            },
            "Перевозка": {
                "type": "object",
                "additionalProperties": False,
                "properties": {
                "Отправитель": {"$ref": "#/$defs/PartyWithIds"},
                "Получатель": {"$ref": "#/$defs/PartyWithIds"},
                "Станция отправления": {"$ref": "#/$defs/StationPlace"},
                "Станция назначения": {"$ref": "#/$defs/StationPlace"},
                "Пограничная станция": {"$ref": "#/$defs/StationPlace"},
                "Условия": {
                    "type": "object",
                    "additionalProperties": False,
                    "properties": {
                    "Условия поставки (Incoterms)": {"type": ["string","null"]}
                    },
                    "required": ["Условия поставки (Incoterms)"]
                },
                "Оплата провозной платы": {
                    "type": "object",
                    "additionalProperties": False,
                    "properties": {
                    "Тип оплаты": {"type": ["string","null"]},
                    "Валюта": {"type": ["string","null"]}
                    },
                    "required": ["Тип оплаты","Валюта"]
                }
                },
                "required": [
                "Отправитель","Получатель",
                "Станция отправления","Станция назначения"
                ]
            },
            "Таможенный пост": {
                "type": "object",
                "additionalProperties": False,
                "properties": {
                "Код ТП": {"type": ["string","null"]},
                "Наименование таможенного поста": {"type": ["string","null"]},
                "Адрес": {"type": ["string","null"]},
                "Номер лицензии (License/Lic.)": {"type": ["string","null"]}
                },
                "required": ["Код ТП","Наименование таможенного поста","Адрес","Номер лицензии (License/Lic.)"]
            },
            "Товары": {
                "type": "array",
                "items": {"$ref": "#/$defs/Item"}
            }
            },
            "required": ["Общая информация","Перевозка","Таможенный пост","Товары"]
        }
        }

    return {"name": "Generic", "schema": {"type": "object"}}

######################################################
def _take_10digits(s: str) -> str:
    if not s: return ""
    m = re.search(r"\b(\d{10})\b", s.replace(" ", ""))
    return m.group(1) if m else ""

def _norm_percent(s: str) -> str:
    if not s:
        return ""
    t = str(s).strip().replace(" ", "")
    t = t.replace(",", ".")
    m = re.search(r"(\d{1,2}(?:\.\d{1,2})?)%", t)
    if not m:
        tl = str(s).lower()
        if "без ндс" in tl or "беспошлин" in tl:
            return "0%"
        return ""
    val = float(m.group(1))
    return f"{int(val)}%" if abs(val - round(val)) < 1e-9 else f"{val}%"

def classify_tnved_gpt(items: list[dict]) -> dict:
    name_map: list[str] = []
    packed: list[dict] = []

    for it in items or []:
        name = (it.get("Наименование") or "").strip()
        extra = (it.get("Дополнительная информация") or "").strip()
        manufacture = (it.get("Производитель") or "").strip()

        full = name
        if extra and extra.lower() != "null":
            full += f" ({extra})"
        if manufacture and manufacture.lower() != "null":
            full += f" — Производитель: {manufacture}"

        if not full:
            continue

        name_map.append(full)
        packed.append({
            "id": len(packed),
            "name": full,
        })

    if not packed:
        return {"items": [], "groups": {}}

    payload = {"items": packed}

    system_prompt = (
        "Ты — эксперт по классификации товаров по ТН ВЭД ЕАЭС и подготовке формулировок для графы 31 ДТ.\n"
        "Если информации недостаточно (назначение/материал/параметры/область применения) — используй web-поиск по типовым описаниям.\n"
        "Не выдумывай бренды/модели, если их нет во входе. Используй только типовые/общедоступные характеристики.\n"
        "Ответ верни СТРОГО одним JSON-объектом без markdown и без пояснений."
    )

    user_prompt = (
        "Сделай за ОДИН ответ:\n"
        "1) Для каждого товара определить 10-значный код ТН ВЭД ЕАЭС и платежи.\n"
        "2) Для каждого УНИКАЛЬНОГО кода ТН ВЭД сформировать Описание группы.\n\n"
        f"Вход:\n{json.dumps(payload, ensure_ascii=False)}\n\n"
        "Верни СТРОГО по схеме (только JSON):\n"
        "{\n"
        '  "items": [\n'
        "    {\n"
        '      "id": number,\n'
        '      "name": string,\n'
        '      "tnved_code": string,\n'
        '      "duty_percent": string,\n'
        '      "vat_percent": string,\n'
        '      "excise_percent": string,\n'
        '      "item_desc31": string\n'
        "    }\n"
        "  ],\n"
        '  "groups": {\n'
        '    "1234567890": { "group_desc": "..." }\n'
        "  }\n"
        "}\n\n"
        "Ограничения:\n"
        "- tnved_code: строго 10 цифр.\n"
        "- item_desc31: короче, чем group_desc, но содержит точные характеристики конкретного товара. 50-100 символов.\n"
        "- group_desc: 100-150 символов, обобщённое описание группы ТНВЭД.\n"
        "- Если сомневаешься — всё равно выбери наилучший код.\n"
    )

    client = gpt_client()
    resp = client.responses.create(
        model="gpt-5",
        tools=[{"type": "web_search"}],
        reasoning={"effort": "medium"},
        input=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ]
    )

    text = (resp.output_text or "").strip()
    raw = text
    if raw.startswith("```"):
        raw = re.sub(r"^```[a-zA-Z0-9]*\s*", "", raw).strip()
        raw = re.sub(r"\s*```$", "", raw).strip()

    try:
        data = json.loads(raw)
    except Exception:
        i = raw.find("{")
        j = raw.rfind("}")
        if i == -1 or j == -1 or j <= i:
            raise RuntimeError("Не найден dict")
        data = json.loads(raw[i:j+1])

    out_items: list[dict] = []
    got_items = data.get("items") or []
    if not isinstance(got_items, list):
        got_items = []

    by_id: dict[int, dict] = {}
    for rec in got_items:
        if not isinstance(rec, dict):
            continue
        try:
            rid = int(rec.get("id"))
        except Exception:
            continue
        by_id[rid] = rec

    for i, full in enumerate(name_map):
        rec = by_id.get(i) or {}
        code = _take_10digits(str(rec.get("tnved_code") or ""))
        duty = _norm_percent(rec.get("duty_percent")) or "0%"
        vat  = _norm_percent(rec.get("vat_percent"))  or "0%"
        exc  = _norm_percent(rec.get("excise_percent")) or "0%"
        desc = (rec.get("item_desc31") or "")

        out_items.append({
            "Наименование": full,
            "Код": code,
            "Пошлина": duty,
            "НДС": vat,
            "Акциз": exc,
            "Техническое описание": str(desc).strip(),
        })

    groups_out: dict[str, dict] = {}
    got_groups = data.get("groups") or {}
    if isinstance(got_groups, dict):
        for k, v in got_groups.items():
            code10 = _take_10digits(str(k))
            if not code10:
                continue
            if isinstance(v, dict):
                gdesc = (v.get("group_desc") or "").strip()
            else:
                gdesc = str(v or "").strip()
            groups_out[code10] = {"Описание группы": gdesc}

    for it in out_items:
        c = (it.get("Код") or "").strip()
        if c and c not in groups_out:
            groups_out[c] = {"Описание группы": ""}

    return {"items": out_items, "groups": groups_out}

def enrich_tnved_if_invoice(parsed: dict, fail_soft: bool = True) -> dict:
    try:
        if not isinstance(parsed, dict):
            return parsed

        if parsed.get("_doc_key") not in {"invoice"}:
            return parsed

        goods = parsed.get("Товары")
        if not isinstance(goods, list) or not goods:
            return parsed

        manufacturer = ""
        try:
            manufacturer = (parsed.get("invoice", {}).get("Отправитель", {}).get("Название компании") or "").strip()
        except Exception:
            manufacturer = ""

        items_for_api: list[dict] = []
        for it in goods:
            if not isinstance(it, dict):
                continue
            name  = (it.get("Наименование") or "").strip()
            extra = (it.get("Дополнительная информация") or "").strip()
            items_for_api.append({
                "Наименование": name,
                "Дополнительная информация": extra,
                "Производитель": manufacturer
            })

        result = classify_tnved_gpt(items_for_api)
        tnved_list = (result or {}).get("items") or []
        groups = (result or {}).get("groups") or {}

        changed = 0
        for i, it in enumerate(goods):
            if not isinstance(it, dict) or i >= len(tnved_list):
                continue

            rec = tnved_list[i] if isinstance(tnved_list[i], dict) else {}

            code = (rec.get("Код") or "").strip()
            duty = (rec.get("Пошлина") or "").strip()
            vat  = (rec.get("НДС") or "").strip()
            exc  = (rec.get("Акциз") or "").strip()
            desc31 = (rec.get("Техническое описание") or "").strip()

            if code:
                it["Код ТНВЭД"] = code
                changed += 1
            if duty:
                it["Пошлина"] = duty
            if vat:
                it["НДС"] = vat
            if exc:
                it["Акциз"] = exc
            if desc31:
                it["Техническое описание"] = desc31

        parsed["_tnved_groups"] = groups if isinstance(groups, dict) else {}

        parsed["_tnved"] = {
            "status": "ok",
            "changed": changed,
            "mode": "overwrite",
            "groups_count": len(parsed["_tnved_groups"]),
            "one_call": True
        }
        return parsed

    except Exception as e:
        if fail_soft:
            parsed["_tnved"] = {"status": "error", "reason": str(e)}
            parsed["_tnved_gpt_error"] = str(e)
            return parsed
        raise RuntimeError(f"TNVED enrichment failed: {e}")

def extract_between_markers(raw: str,
                            start_tag: str = "<#START_JSON#>",
                            end_tag: str   = "<#END_JSON#>") -> Optional[str]:
    s = raw.rfind(start_tag)
    if s == -1:
        return None
    s += len(start_tag)
    e = raw.find(end_tag, s)
    if e == -1:
        return None

    chunk = raw[s:e].strip()
    m = _CODEBLOCK_RE.match(chunk)
    if m:
        chunk = m.group(1).strip()
    if (chunk.startswith('"') and chunk.endswith('"')) or (chunk.startswith("'") and chunk.endswith("'")):
        try:
            chunk = json.loads(chunk)
        except Exception:
            pass

    m = _CODEBLOCK_RE.match(chunk)
    if m:
        chunk = m.group(1).strip()

    return chunk

_CODEBLOCK_RE = re.compile(r"^```(?:json)?\s*(.*?)\s*```$", re.DOTALL)

def extract_json_fallback(raw: str) -> Optional[Dict[str, Any]]:
    try:
        return json.loads(raw.strip())
    except Exception:
        pass
    m = re.search(r"```json\s*(\{.*?\})\s*```", raw, flags=re.DOTALL)
    if m:
        try: return json.loads(m.group(1))
        except Exception: pass

    depth=0; start=-1; in_str=False; esc=False
    for i,ch in enumerate(raw):
        if esc: esc=False; continue
        if ch=='\\' and in_str: esc=True; continue
        if ch=='"': in_str = not in_str; continue
        if in_str: continue
        if ch=='{':
            if depth==0: start=i
            depth+=1
        elif ch=='}':
            if depth>0:
                depth-=1
                if depth==0 and start!=-1:
                    try: return json.loads(raw[start:i+1])
                    except Exception: break

    s,e = raw.find("{"), raw.rfind("}")
    if s!=-1 and e!=-1 and e>s:
        try: return json.loads(raw[s:e+1])
        except Exception: pass
    return None

def yandex_client() -> OpenAI:
    if not (YANDEX_API_KEY and YANDEX_FOLDER):
        raise RuntimeError("YANDEX_API_KEY или YANDEX_FOLDER_ID не заданы")
    return OpenAI(
        api_key=YANDEX_API_KEY,
        base_url="https://rest-assistant.api.cloud.yandex.net/v1",
        project=YANDEX_FOLDER,
    )

def extract_docx_text_with_meta(file_bytes: bytes, filename: str = "") -> tuple[str, dict]:
    buf = io.BytesIO(file_bytes)
    doc = Document(buf)

    parts: list[str] = []

    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)

    for table in doc.tables:
        for row in table.rows:
            cells = [(cell.text or "").strip() for cell in row.cells]
            line = " | ".join([c for c in cells if c])
            if line:
                parts.append(line)

    text = "\n".join(parts)

    meta = {
        "engine": "python-docx",
        "filename": filename,
        "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "chars": len(text),
    }
    return text, meta


def extract_xlsx_text_with_meta(file_bytes: bytes, filename: str = "") -> tuple[str, dict]:
    wb = load_workbook(io.BytesIO(file_bytes), data_only=True)

    parts: list[str] = []
    sheets_used = 0
    cells_used = 0

    for ws in wb.worksheets:
        sheets_used += 1
        parts.append(f"--- SHEET: {ws.title} ---")

        max_r = ws.max_row or 0
        max_c = ws.max_column or 0

        for r in range(1, max_r + 1):
            row_vals: list[str] = []
            empty_row = True

            for c in range(1, max_c + 1):
                v = ws.cell(row=r, column=c).value
                if v is None:
                    row_vals.append("")
                    continue
                s = str(v).strip()
                if s:
                    empty_row = False
                row_vals.append(s)
                cells_used += 1

            if not empty_row:
                line = " | ".join([x for x in row_vals if x != ""])
                if line:
                    parts.append(line)

        parts.append("")

    text = "\n".join(parts).strip()

    meta = {
        "engine": "openpyxl",
        "filename": filename,
        "mime": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "chars": len(text),
        "sheets_used": sheets_used,
        "cells_used": cells_used,
    }
    return text, meta

def call_yandexgpt(file_bytes: bytes, filename: str, doc_key: str, mime: Optional[str] = None) -> Optional[Dict[str, Any]]:

    client = yandex_client()

    schema_meta = json_schema_for(doc_key)  # {"name": "...", "schema": {...}}
    schema_json = json.dumps(schema_meta.get("schema", {}), ensure_ascii=False)
    fn_lower = (filename or "").lower()
    mime = mime or ""

    is_docx = (
        fn_lower.endswith(".docx")
        or "wordprocessingml" in mime  
    )

    is_xlsx = (
        fn_lower.endswith(".xlsx")
        or "spreadsheetml" in mime
        or mime in {
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "application/vnd.ms-excel",
        }
    )

    if is_docx:
        plain_text, ocr_meta = extract_docx_text_with_meta(file_bytes, filename)
    elif is_xlsx:
        plain_text, ocr_meta = extract_xlsx_text_with_meta(file_bytes, filename)
    else:
        plain_text, ocr_meta = extract_text_with_meta(
            file_bytes,
            mime or "application/octet-stream",
        )

    if not plain_text or not plain_text.strip():
        raise RuntimeError("Не удалось извлечь текст из файла")
    
    user_prompt = build_prompt(doc_key, filename, plain_text)
    sys_text = (
        "Направляю тебе json-schema для ответа.\n"
        "Выведи строго один JSON-объект между тегами <#START_JSON#> и <#END_JSON#>.\n"
        "Никакого текста до/после этих тегов.\n"
        "Если значения не найдены — ставь null.\n"
        "JSON должен соответствовать следующей JSON Schema:\n"
        f"{schema_json}")
    
    response = client.responses.create(
        model=f"gpt://{YANDEX_CLOUD_FOLDER}/{YANDEX_CLOUD_MODEL}",
        temperature=0.0,
        instructions=f"{sys_text}",
        input=f"{user_prompt}",
        max_output_tokens=10000
    )
    raw = (response.output_text or "").strip()

    parse_meta = {}                     
    parsed = None

    chunk = extract_between_markers(raw)
    if chunk is not None:
        try:
            parsed = json.loads(chunk)
        except Exception as e1:
            try:
                inner = json.loads(chunk)
                if isinstance(inner, str):
                    parsed = json.loads(inner)
            except Exception as e2:
                parse_meta = {"source": "markers", "error": f"{e1}; nested: {e2}"}
                parsed = None
    else:
        parse_meta = {"source": "markers", "error": "not found"}

    if parsed is None:
        fb = extract_json_fallback(raw)
        if fb is not None:
            parsed = fb
            parse_meta = {"source": "fallback"}
        else:
            return {
                "raw": raw,
                "_error": "json_parse_failed",
                "_parse_meta": parse_meta,
                "_schema_name": schema_meta.get("name"),
                "_ocr": ocr_meta,
                "_source_text": plain_text,
            }

    return parsed

################## FastAPI docs ##################
class EnqueueBody(BaseModel):
    decl_id: int
    file_id: int
    doc_key: str

class EnqueueResp(BaseModel):
    job_id: int
    status: str = "queued"

class JobResp(BaseModel):
    job_id: int
    status: str
    error: Optional[str] = None
    result: Optional[Dict[str, Any]] = None

class OCRDebugResp(BaseModel):
    mime: str
    chars: int
    preview: str
    truncated: bool

################## Worker ##################

_stop = threading.Event()

@asynccontextmanager
async def lifespan(app: FastAPI):
    threading.Thread(target=worker_loop, daemon=True).start()
    yield
    _stop.set()

app = FastAPI(title="AI-Декларант API", version="1.0", lifespan=lifespan)

auth_router = APIRouter(prefix="/auth", tags=["auth"])
users_router = APIRouter(prefix="/users", tags=["users"])
decl_router = APIRouter(prefix="/v1/declarations", tags=["declarations"])
jobs_router = APIRouter(prefix="/v1", tags=["jobs"])
admin_router = APIRouter(prefix="/admin", tags=["admin"])

misc_router = APIRouter(tags=["misc"])
tnved_router = APIRouter(tags=["tnved"])
company_router = APIRouter(prefix="/company", tags=["company"])
debug_router = APIRouter(prefix="/debug", tags=["debug"])
graphs_router = APIRouter(prefix="/v1/graphs", tags=["graphs"])
files_router = APIRouter(tags=["files"])






origins = [
    "https://ai-declar.ru",
    "http://ai-declar.ru",
    "https://declarant-ai.ru",
    "http://declarant-ai.ru",
    "http://localhost:5173",
    "http://127.0.0.1:5173"
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.exception_handler(StarletteHTTPException)
async def http_exception_handler(request: Request, exc: StarletteHTTPException):
    return PlainTextResponse(
        content=str(exc.detail),
        status_code=exc.status_code,
    )

def get_current_user(request: Request) -> Dict[str, Any]:
    auth = request.headers.get("Authorization") or ""
    if not auth.startswith("Bearer "):
        raise HTTPException(401, "Missing Bearer token")

    token = auth.replace("Bearer ", "", 1).strip()
    try:
        data = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALG])
        sub = data.get("sub")
        if not sub:
            raise HTTPException(401, "Invalid token")
        user_id = int(sub)
    except JWTError:
        raise HTTPException(401, "Invalid token")

    user = get_user_by_id(user_id)
    if not user:
        raise HTTPException(401, "User not found")
    return user

@jobs_router.post("/jobs", response_model=EnqueueResp)
def enqueue_job(body: EnqueueBody, current=Depends(get_current_user)):
    owner_id = int(get_declaration_user_id(int(body.decl_id)) or 0)
    if owner_id != int(current["id"]):
        raise HTTPException(403, "Forbidden (not your declaration)")
    jid = jobs_create(body.decl_id, body.file_id, body.doc_key)
    return EnqueueResp(job_id=jid)

@jobs_router.get("/jobs/{job_id}", response_model=JobResp)
def job_status(job_id: int):
    row = jobs_get(job_id)
    if not row:
        raise HTTPException(404, "job not found")
    return JobResp(
        job_id=row["id"],
        status=row["status"],
        error=row.get("error_text"),
        result=row.get("result_json")
    )

@debug_router.post("/ocr")
async def debug_ocr(file: UploadFile = File(...)):
    file_bytes = await file.read()
    text, meta = extract_text_with_meta(
        file_bytes=file_bytes,
        mime=file.content_type or ""
    )

    return {
        "meta": meta,
        "text_preview": text[:4000], 
        "text_length": len(text or ""),
    }


@decl_router.get("/{decl_id}/jobs")
def jobs_by_decl(decl_id: int):
    return jobs_list_by_decl(decl_id)

def persist_doc_json(decl_id: int, user_id: int, doc_key: str, data: dict):
    dk = f"{doc_key}_json"

    payload = json.dumps(data, ensure_ascii=False, indent=2).encode("utf-8")
    filename = f"{dk}.json"
    linked = list_declaration_files(decl_id) or []
    for r in linked:
        if str(r.get("doc_key") or "") == dk:
            try:
                unlink_file_from_declaration(decl_id, r["file_id"])
            except Exception:
                pass

    new_file_id = add_file(user_id, filename, "application/json", payload)
    link_file_to_declaration(decl_id, new_file_id, dk)

class TnvedGoodsItemOut(BaseModel):
    index: int
    name: str
    manufacturer: Optional[str] = None
    extra_info: Optional[str] = None


class TnvedGoodsListOut(BaseModel):
    goods: List[TnvedGoodsItemOut]
    count: int


class TnvedEnrichItem(BaseModel):
    index: int
    extra_info: Optional[str] = None


class TnvedEnrichRequest(BaseModel):
    items: List[TnvedEnrichItem]

@decl_router.get("/{decl_id}/tnved/goods", tags=["tnved"], response_model=TnvedGoodsListOut,)
def api_get_tnved_goods(decl_id: int):
    parsed = get_declaration_invoice_json(decl_id)
    if not parsed:
        raise HTTPException(
            status_code=404,
            detail="invoice_json не найден для этой декларации",
        )

    goods = parsed.get("Товары")
    if not isinstance(goods, list) or not goods:
        raise HTTPException(
            status_code=400,
            detail="В invoice_json нет списка 'Товары'",
        )

    sender = (
        parsed.get("Отправитель")                         
        or (parsed.get("invoice") or {}).get("Отправитель")  
        or {}
    )

    default_manufacturer = (
        sender.get("Название компании")
        or sender.get("Наименование компании")
        or sender.get("Company name")
        or ""
    ).strip()

    out: List[TnvedGoodsItemOut] = []

    for idx, g in enumerate(goods):
        if not isinstance(g, dict):
            continue

        name = (
            g.get("Наименование")
            or g.get("Описание")
            or g.get("Name")
            or ""
        )
        name = str(name).strip()

        manufacturer = (
            g.get("Производитель")
            or g.get("Производитель товара")
            or default_manufacturer
            or ""
        )
        manufacturer = str(manufacturer).strip() or None

        extra = (
            g.get("_user_extra_info")
            or g.get("Дополнительная информация")
            or ""
        ).strip() or None

        out.append(
            TnvedGoodsItemOut(
                index=idx,
                name=name,
                manufacturer=manufacturer,
                extra_info=extra,
            )
        )

    return TnvedGoodsListOut(goods=out, count=len(out))

@decl_router.post("/{decl_id}/tnved/enrich", tags=["tnved"])
def api_enrich_tnved_for_invoice(decl_id: int, body: TnvedEnrichRequest,):
    parsed = get_declaration_invoice_json(decl_id)
    if not parsed:
        raise HTTPException(
            status_code=404,
            detail="invoice_json не найден для этой декларации",
        )

    goods = parsed.get("Товары")
    if not isinstance(goods, list) or not goods:
        raise HTTPException(
            status_code=400,
            detail="В invoice_json нет списка 'Товары'",
        )

    extra_by_index: Dict[int, str] = {}
    for it in body.items or []:
        try:
            idx = int(it.index)
        except Exception:
            continue
        val = (it.extra_info or "").strip()
        extra_by_index[idx] = val

    for idx, g in enumerate(goods):
        if not isinstance(g, dict):
            continue
        if idx not in extra_by_index:
            continue
        val = extra_by_index[idx]
        g["Дополнительная информация"] = val
        g["_user_extra_info"] = val

    parsed.setdefault("_doc_key", "invoice")
    try:
        parsed_after = enrich_tnved_if_invoice(parsed, fail_soft=False)
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Ошибка при определении ТН ВЭД: {e}",
        )
    save_declaration_invoice_json(decl_id, parsed_after)
    return {
        "ok": True,
        "updated_items": len(extra_by_index),
        "tnved": parsed_after.get("_tnved", {}),
    }

def worker_loop():
    while not _stop.is_set():
        job = jobs_claim_next()
        if not job:
            time.sleep(0.8)
            continue
        jid, fid, doc_key = job["id"], job["file_id"], job["doc_key"]
        try:
            rec = get_file(fid)
            blob, fname = rec["file_data"], rec["filename"]
            mime = rec.get("mime")
            result = call_yandexgpt(blob, fname, doc_key, mime=mime)  

            jobs_finish_ok(jid, result)
            user_id = rec.get("user_id") 
            persist_doc_json(job["decl_id"], user_id, doc_key, result)  
            time.sleep(1)

        except Exception:
            jobs_finish_err(jid, traceback.format_exc())




############### ЕГРЮЛ API OFDATA ###############
def parse_ofdata_company(payload: Dict[str, Any]) -> Dict[str, Any]:
    if not payload:
        raise ValueError("Пустой ответ OfData")


    inn = payload.get("ИНН") or ""
    kpp = payload.get("КПП") or ""
    ogrn = payload.get("ОГРН") or ""
    full_name = payload.get("НаимПолн") or ""
    short_name = payload.get("НаимСокр") or ""

    addr_block = (payload.get("ЮрАдрес") or {}).get("АдресРФ") or ""
    city = (payload.get("ЮрАдрес") or {}).get("НасПункт") or ""

    index = extract_index(addr_block) if addr_block else ""

    region = ""
    street = None
    house = None

    if addr_block:
        parts = [p.strip() for p in addr_block.split(",") if p.strip()]

        if len(parts) > 1:
            if parts[1] == city:
                region = ""
            else:
                region = parts[1]

        street_index = None
        for i, p in enumerate(parts):
            lp = p.lower()
            if (
                "ул." in lp
                or "улица" in lp
                or "пр-кт" in lp
                or "просп" in lp
                or "ш." in lp
                or "шоссе" in lp
                or "набер" in lp
            ):
                street = p
                street_index = i
                break

        if street_index is not None and street_index + 1 < len(parts):
            house = ", ".join(parts[street_index + 1 :])

    return {
        "inn": inn,
        "kpp": kpp,
        "ogrn": ogrn,
        "full_name": full_name,
        "short_name": short_name,
        "index": index,
        "region": region,
        "city": city,
        "street": street,
        "house": house,
    }

@company_router.get("/ofdata")
def get_company_ofdata(inn: str = Query(..., min_length=10, max_length=12),):
    if not inn.isdigit():
        raise HTTPException(status_code=400, detail="ИНН должен состоять только из цифр")

    if OFDATA_API_KEY is None:
        raise HTTPException(status_code=500, detail="OFDATA_API_KEY не настроен на сервере")

    try:
        payload = {
            "key": OFDATA_API_KEY,
            "inn": inn,
        }
        r = httpx.post(OFDATA_URL, json=payload, timeout=5.0)
    except httpx.RequestError as e:
        raise HTTPException(status_code=502, detail=f"Ошибка запроса к OfData: {e}") from e

    if r.status_code == 404:
        raise HTTPException(status_code=404, detail="Компания с таким ИНН не найдена в ЕГРЮЛ")

    try:
        data = r.json()
    except ValueError:
        raise HTTPException(status_code=502, detail="Некорректный JSON от OfData")

    meta = data.get("meta") or {}
    left = None
    try:
        left = 100 - int(meta.get("today_request_count", 0))
    except Exception:
        left = None

    company_raw = data.get("data") or {}
    try:
        company = parse_ofdata_company(company_raw)
    except ValueError as e:
        raise HTTPException(status_code=502, detail=str(e))

    return {
        "limits_left": left,
        "company": company,
        "raw": company_raw, 
    }

############### TNVED ###############

class DetectIn(BaseModel):
    manufacturer: str
    product: str
    extra: Optional[str] = None
    user_id: Optional[int] = None

class AltItem(BaseModel):
    code: Optional[str] = None
    reason: Optional[str] = None

class Payments(BaseModel):
    duty: Optional[str] = None
    vat: Optional[str] = None
    excise: Optional[str] = None
    fees: Optional[str] = None

class DetectOut(BaseModel):
    code: str
    duty: str
    vat: str
    raw: Optional[str] = None
    description: Optional[str] = None
    tech31: Optional[str] = None
    decl31: Optional[str] = None
    classification_reason: Optional[str] = None
    alternatives: Optional[List[AltItem]] = None
    payments: Optional[Payments] = None
    requirements: Optional[List[str]] = None

class FeedbackIn(BaseModel):
    acc_code: int
    desc_31: int
    reason_clarity: int
    ui: int
    comment: Optional[str] = None

    manufacturer: Optional[str] = None
    product: Optional[str] = None
    extra: Optional[str] = None

    code: Optional[str] = None
    tech31: Optional[str] = None
    decl31: Optional[str] = None

def _extract_json_block(text: str) -> Optional[Dict[str, Any]]:
    if not text:
        return None
    try:
        start = text.index("{")
        end = text.rfind("}")
        blob = text[start:end+1]
        return json.loads(blob)
    except Exception:
        return None

def _stringify_tech31(val) -> str:
    if val is None:
        return ""
    if isinstance(val, str):
        return val.strip()
    if isinstance(val, dict):
        parts = []
        for k, v in val.items():
            k_s = str(k).strip().capitalize()
            if isinstance(v, (list, tuple)):
                v_s = "; ".join(str(x).strip() for x in v if str(x).strip())
            elif isinstance(v, dict):
                v_s = "; ".join(f"{kk}: {vv}" for kk, vv in v.items())
            else:
                v_s = str(v).strip()
            if v_s:
                parts.append(f"- {k_s}: {v_s}")
        return "\n".join(parts)
    if isinstance(val, (list, tuple)):
        return "\n".join(f"- {str(x).strip()}" for x in val if str(x).strip())
    return str(val).strip()

def _normalize_alternatives(val):
    out = []
    if isinstance(val, dict):
        for k, v in val.items():
            out.append({"code": str(k), "reason": str(v)})
    elif isinstance(val, (list, tuple)):
        for it in val:
            if isinstance(it, dict):
                out.append({
                    "code": str(it.get("code", "") or it.get("код", "") or ""),
                    "reason": str(it.get("reason", "") or it.get("обоснование", "") or "")
                })
            else:
                out.append({"code": str(it), "reason": ""})
    elif val:
        out.append({"code": str(val), "reason": ""})
    return out

def _normalize_payments(val, fallback_duty: str, fallback_vat: str):
    d = {"duty": fallback_duty, "vat": fallback_vat, "excise": "—", "fees": "—"}
    if isinstance(val, dict):
        for k in ("duty", "vat", "excise", "fees"):
            if k in val and val[k] is not None:
                d[k] = str(val[k]).strip()
    return d

def save_feedback_to_db(fb: FeedbackIn, request: Request) -> None:
    conn = get_conn()
    if conn is None:
        return

    try:
        with conn: 
            with conn.cursor() as cur:
                cur.execute(
                    """
                    INSERT INTO feedback (
                        acc_code_rating,
                        tech31_rating,
                        reason_rating,
                        ui_rating,
                        req_manufacturer,
                        req_product,
                        req_extra,
                        res_code,
                        res_tech31,
                        res_decl31,
                        comment,
                        user_agent,
                        client_ip
                    )
                    VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)
                    """,
                    (
                        fb.acc_code,
                        fb.desc_31,
                        fb.reason_clarity,
                        fb.ui,
                        fb.manufacturer or "",
                        fb.product or "",
                        fb.extra or "",
                        fb.code or "",
                        fb.tech31 or "",
                        fb.decl31 or "",
                        fb.comment or "",
                        request.headers.get("user-agent", ""),
                        request.client.host if request.client else None,
                    ),
                )
    finally:
        conn.close()

def _normalize_requirements(val):
    if isinstance(val, (list, tuple)):
        return [str(x).strip() for x in val if str(x).strip()]
    if isinstance(val, str):
        import re as _re
        items = [s.strip(" -•\t") for s in _re.split(r"[\n;]+", val) if s.strip()]
        return items or [val.strip()]
    if val:
        return [str(val)]
    return []

@tnved_router.post("/tnved/detect", response_model=DetectOut)
def detect(inp: DetectIn, request: Request):
    ip = client_ip(request)
    rate_limit(f"tnved:{ip}", limit=1, window_sec=60)
    t0 = time.time()
    if inp.user_id is not None:
        uid = int(inp.user_id)
        if uid > 0:
            try:
                credits_spend(
                    user_id=uid,
                    amount=1,
                    reason="use",
                    ref_type="tnved_detect",
                    ref_id=None,
                    meta={
                        "path": "/tnved/detect",
                        "ip": request.client.host if request.client else None,
                        "ua": request.headers.get("user-agent", ""),
                    },
                )
            except ValueError as e:
                # 402 — “недостаточно запросов / пакет истёк”
                raise HTTPException(status_code=402, detail=str(e))
    full = (inp.product or "").strip()
    if inp.extra and inp.extra.strip().lower() != "null":
        full += f" ({inp.extra.strip()})"
    if inp.manufacturer and inp.manufacturer.strip().lower() != "null":
        full += f" — Производитель: {inp.manufacturer.strip()}"

    if not full:
        raise HTTPException(status_code=400, detail="Поля пустые")

    system_msg = """Ты — эксперт по классификации товаров по ТН ВЭД ЕАЭС и по подготовке текстов для графы 31 декларации на товары.
    Ты — эксперт по классификации товаров по ТН ВЭД ЕАЭС и по подготовке текстов для графы 31 декларации на товары.
    Твоя задача: по краткому описанию товара определить наиболее вероятный 10-значный код ТН ВЭД ЕАЭС, указать ставки платежей и сформировать подробное техническое описание товара.
    Если предоставленной информации недостаточно для уверенной классификации (нет назначения, материалов, электрических параметров, области применения и т.п.), ты должен сначала получить недостающие сведения через web-поиск по типовым описаниям схожих товаров и уже на основе найденного сформировать итоговое описание. Используй только общедоступные и типовые характеристики, не выдумывай конкретные модели и бренды, если их нет во входных данных. Делай оговорки: «по типовым техническим характеристикам для такого вида товара».
    Результат верни строго в виде одного json-объекта
    Структура JSON (поля на русском):
    
    {
      "code": "10-значный код или \"UNKNOWN\"",
      "duty": "проценты или \"UNKNOWN\"",
      "vat": "проценты или \"UNKNOWN\"",
      "tech31": "подробное структурированное техописание: 1) назначение; 2) конструкция и материалы; 3) основные технические/электрические параметры (если применимо); 4) условия эксплуатации; 5) комплектация. Объем не меньше 100 слов. Если часть данных взята из типовых открытых источников — так и укажи.",
      "decl31": "готовая формулировка для графы 31 декларации на товары, краткая, без лишних пояснений, в одном абзаце, с указанием основных отличительных признаков и назначения. Без слов «примерно», «возможно», «как правило».",
      "classification_reason": "обоснование выбора позиции ТН ВЭД (ОПИ, примечания к группе/товарной позиции, признаки товара). Если есть неопределенность — укажи диапазон возможных кодов и чего не хватает.",
      "alternatives": [
        {"code": "возможный_код", "reason": "в каких случаях применим"}
      ],
      "payments": {
        "duty": "% или \"UNKNOWN\"",
        "vat": "% или \"UNKNOWN\"",
        "excise": "— или значение",
        "fees": "— или значение"
      },
      "requirements": [
        "ТР ЕАЭС, безопасность, лицензирование, сертификация — если применимо"
      ]
    }
    
    Требования:
    - не добавляй никаких комментариев вне JSON;
    - не меняй имена полей;
    - если веб-поиск не дал точных параметров — пиши «по типовым характеристикам для данного вида товара».
    
    """
    user_msg = (
        "Определи 10-значный код ТН ВЭД для товара и верни результат СТРОГО в виде JSON.\n"
        "Вход:\n"
        f"{json.dumps({'Наименование': full}, ensure_ascii=False)}"
    )

    try:
        client = gpt_client()
        resp = client.responses.create(
            model="gpt-5",
            tools=[{"type": "web_search"}],
            reasoning={"effort": "medium"},
            input=[
                {"role": "system", "content": system_msg},
                {"role": "user", "content": user_msg},
            ],
        )
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Ошибка GPT API: {e}")

    text = (resp.output_text or "").strip()
    data = _extract_json_block(text) or {}

    code = (data.get("code") or "").strip()
    if not _take_10digits(code):
        guessed = _take_10digits(text)
        code = guessed or (code if code.upper().startswith("UNKNOWN") else "")

    duty = _norm_percent(data.get("duty") or "")
    vat  = _norm_percent(data.get("vat") or "")

    code = code or "UNKNOWN"
    duty = duty or "UNKNOWN"
    vat  = vat or "UNKNOWN"

    tech31 = _stringify_tech31(data.get("tech31"))
    alternatives = _normalize_alternatives(data.get("alternatives"))
    payments = _normalize_payments(data.get("payments"), fallback_duty=duty, fallback_vat=vat)
    requirements = _normalize_requirements(data.get("requirements"))
    decl31 = (data.get("decl31") or "").strip()


    latency_ms = int((time.time() - t0) * 1000)
    if inp.user_id:
        try:
            tnved_requests_add(
                user_id=int(inp.user_id),
                query_text=full,
                input_json={"manufacturer": inp.manufacturer, "product": inp.product, "extra": inp.extra},
                result_json=out.model_dump(),
                status="done",
                credits_spent=1,
                model="gpt-5",
                latency_ms=latency_ms,
            )
        except Exception as e:
            print("[tnved_requests_add] error:", e)

    out = DetectOut(
        code=code,
        duty=duty,
        vat=vat,
        raw=text,
        description=(data.get("description") or ""),
        tech31=tech31,
        decl31=decl31,
        classification_reason=(data.get("classification_reason") or ""),
        alternatives=alternatives,
        payments=payments,
        requirements=requirements,
    )
    return out


@misc_router.post("/feedback", tags=["feedback"])
def feedback(fb: FeedbackIn, request: Request):
    saved = False
    error = None

    try:
        save_feedback_to_db(fb, request)
        saved = True
    except Exception as e:
        error = str(e)
        print(f"[feedback] DB error: {e}")

    return {
        "status": "ok",
        "saved": saved,  
        "error": error  
    }

@misc_router.get("/")
def root():
    return {"status": "Проверка работоспособности"}

############### AUTH / USERS / PROFILE ###############

class UserOut(BaseModel):
    id: int
    email: EmailStr
    name: str
    surname: str
    avatar_path: str

class UserRegisterIn(BaseModel):
    name: str
    surname: str
    email: EmailStr
    password: str

class UserLoginIn(BaseModel):
    email: EmailStr
    password: str

class UserProfileIn(BaseModel):
    name: Optional[str] = None
    surname: Optional[str] = None
    position: Optional[str] = None
    phone: Optional[str] = None
    email: Optional[EmailStr] = None
    company: Optional[str] = None
    address: Optional[str] = None
    notes: Optional[str] = None
    avatar_path: Optional[str] = None

class AvatarUploadResp(BaseModel):
    file_id: int
    avatar_path: str 

class DeclarationOut(BaseModel):
    id: int
    title: str
    created_at: datetime
    attached_file_id: Optional[int] = None
    file_name: Optional[str] = None

class DeclarationCreateIn(BaseModel):
    title: str
    created_date: Optional[date] = None  

class DeclFileOut(BaseModel):
    link_id: int
    doc_key: str
    created_at: datetime
    file_id: int
    filename: str
    mime: Optional[str]
    size_bytes: int

class FileUploadResp(BaseModel):
    file_id: int
    decl_id: int
    doc_key: str
    filename: str
    size_bytes: int

APP_TZ = ZoneInfo("Europe/Moscow") 


pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")

def hash_password(password: str) -> str:
    return pwd_context.hash(password)

def verify_password(password: str, password_hash: str) -> bool:
    try:
        return pwd_context.verify(password, password_hash)
    except Exception:
        return False

def looks_like_hash(s: str) -> bool:
    if not s:
        return False
    return s.startswith("$2a$") or s.startswith("$2b$") or s.startswith("$2y$")

def _user_to_out(row: Dict[str, Any]) -> UserOut:
    return UserOut(
        id=row["id"],
        email=row["email"],
        name=row.get("name") or "",
        surname=row.get("surname") or "",
        avatar_path=row.get("avatar_path") or "",
    )

class TokenOut(BaseModel):
    access_token: str
    token_type: str = "bearer"
    user: UserOut

def create_access_token(*, user_id: int) -> str:
    exp = datetime.utcnow() + timedelta(minutes=ACCESS_TOKEN_EXPIRE_MIN)
    payload = {"sub": str(user_id), "exp": exp}
    return jwt.encode(payload, JWT_SECRET, algorithm=JWT_ALG)



    token = auth.replace("Bearer ", "", 1).strip()
    try:
        data = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALG])
        sub = data.get("sub")
        if not sub:
            raise HTTPException(401, "Invalid token")
        user_id = int(sub)
    except JWTError:
        raise HTTPException(401, "Invalid token")

    user = get_user_by_id(user_id)
    if not user:
        raise HTTPException(401, "User not found")
    return user

def require_admin(user: Dict[str, Any] = Depends(get_current_user)) -> Dict[str, Any]:
    if (user.get("role") or "user") != "admin":
        raise HTTPException(403, "Admin only")
    return user

def require_owner(path_user_id: int, user: Optional[Dict[str, Any]]) -> None:
    if not user:
        raise HTTPException(401, "Not authenticated")
    if int(path_user_id) != int(user.get("id")):
        raise HTTPException(403, "Forbidden (not your resource)")


@auth_router.post("/register", response_model=TokenOut)
def auth_register(body: UserRegisterIn):
    existing = get_user_by_email(body.email)
    if existing:
        raise HTTPException(400, "Пользователь с таким email уже существует")

    # пока поддерживаем миграцию: храним bcrypt
    user_id = create_user(
        name=body.name,
        surname=body.surname,
        email=body.email,
        password=hash_password(body.password),
    )
    user = get_user_by_id(user_id)

    token = create_access_token(user_id=int(user_id))
    return TokenOut(access_token=token, user=_user_to_out(user))

# @app.post("/auth/login", response_model=UserOut)
# def auth_login(body: UserLoginIn):
#     user = get_user_by_email(body.email)
#     if not user:
#         raise HTTPException(401, "Неверный email или пароль")

#     stored = (user.get("password") or "").strip()
#     ok = False

#     if looks_like_hash(stored):
#         ok = verify_password(body.password, stored)
#     else:
#         ok = (stored == body.password)
#         if ok:
#             try:
#                 update_user(user["id"], password=hash_password(body.password))
#             except Exception:
#                 pass

#     if not ok:
#         raise HTTPException(401, "Неверный email или пароль")

#     return _user_to_out(user)

@auth_router.post("/login", response_model=TokenOut)
def auth_login(body: UserLoginIn, request: Request):
    ip = client_ip(request)
    rate_limit(f"login:{ip}:{body.email}", limit=10, window_sec=60)
    user = get_user_by_email(body.email)
    if not user:
        raise HTTPException(401, "Неверный email или пароль")

    stored = (user.get("password") or "").strip()

    ok = False
    if looks_like_hash(stored):
        ok = verify_password(body.password, stored)
    else:
        ok = (stored == body.password)
        if ok:
            try:
                update_user(int(user["id"]), password=hash_password(body.password))
            except Exception:
                pass

    if not ok:
        raise HTTPException(401, "Неверный email или пароль")

    token = create_access_token(user_id=int(user["id"]))
    return TokenOut(access_token=token, user=_user_to_out(user))

@users_router.get("/{user_id}/profile", response_model=UserProfileIn)
def get_profile(user_id: int, current=Depends(get_current_user)):
    require_owner(user_id, current)
    user = get_user_by_id(user_id)
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    return UserProfileIn(
        name=user.get("name"),
        surname=user.get("surname"),
        position=user.get("position"),
        phone=user.get("phone"),
        email=user.get("email"),
        company=user.get("company"),
        address=user.get("address"),
        notes=user.get("notes"),
        avatar_path=user.get("avatar_path"),
    )

@users_router.put("/{user_id}/profile", response_model=UserProfileIn)
def update_profile(user_id: int, body: UserProfileIn, current=Depends(get_current_user)):
    require_owner(user_id, current)
    user = get_user_by_id(user_id)
    if not user:
        raise HTTPException(status_code=404, detail="User not found")

    fields = body.model_dump(exclude_unset=True)
    if fields:
        update_user(user_id, **fields)

    updated = get_user_by_id(user_id)
    return UserProfileIn(
        name=updated.get("name"),
        surname=updated.get("surname"),
        position=updated.get("position"),
        phone=updated.get("phone"),
        email=updated.get("email"),
        company=updated.get("company"),
        address=updated.get("address"),
        notes=updated.get("notes"),
        avatar_path=updated.get("avatar_path"),
    )

@users_router.get("/{user_id}/declarations", response_model=list[DeclarationOut])
def api_list_declarations(user_id: int, current=Depends(get_current_user)):
    require_owner(user_id, current)
    rows = list_declarations(user_id, limit=500) or []
    return [DeclarationOut(**r) for r in rows]

@users_router.post("/{user_id}/declarations", response_model=DeclarationOut)
def api_create_declaration(user_id: int, body: DeclarationCreateIn, current=Depends(get_current_user)):
    require_owner(user_id, current)
    now = datetime.now(APP_TZ)

    created_at = None
    if body.created_date:
        created_at = now.replace(
            year=body.created_date.year,
            month=body.created_date.month,
            day=body.created_date.day,
        )

    decl_id = add_declaration(user_id=user_id, title=body.title, created_at=created_at)
    row = next((d for d in list_declarations(user_id) if d["id"] == decl_id), None)
    if not row:
        raise HTTPException(500, "Не удалось прочитать созданную декларацию")
    return DeclarationOut(**row)


@files_router.get("/declarations/{decl_id}/files", response_model=list[DeclFileOut])
def api_list_decl_files(decl_id: int):
    rows = list_declaration_files(decl_id) or []
    return [DeclFileOut(**r) for r in rows]

@files_router.post("/declarations/{decl_id}/files", response_model=FileUploadResp)
async def api_upload_decl_file(decl_id: int,user_id: int = Form(...),doc_key: str = Form(...),file: UploadFile = File(...), current=Depends(get_current_user)):
    owner_id = int(get_declaration_user_id(decl_id) or 0)
    if owner_id != int(current["id"]):
        raise HTTPException(403, "Forbidden (not your declaration)")
    
    data = await file.read()
    pages = 0
    mime = file.content_type or ""

    if mime == "application/pdf":
        try:
            pages = pdf_page_count(data)
        except Exception:
            pages = 0

    MAX_PDF_PAGES = 3

    if pages > MAX_PDF_PAGES:
        raise HTTPException(
            status_code=413,
            detail=(
                f"PDF содержит {pages} страниц. "
                f"Допустимо не более {MAX_PDF_PAGES} страниц для автоматического OCR. "
                f"Разбейте документ на части по ссылке ниже:"
            )
        )
    mime = file.content_type or "application/octet-stream"
    file_id = add_file(int(current["id"]), file.filename, mime, data)
    link_file_to_declaration(decl_id, file_id, doc_key, replace=False)

    return FileUploadResp(
        file_id=file_id,
        decl_id=decl_id,
        doc_key=doc_key,
        filename=file.filename,
        size_bytes=len(data),
    )

@files_router.get("/files/{file_id}/download")
def api_download_file(file_id: int, current=Depends(get_current_user)):
    rec = get_file(file_id)
    if not rec:
        raise HTTPException(status_code=404, detail="Файл не найден")

    if int(rec.get("user_id") or 0) != int(current["id"]) and (current.get("role") != "admin"):
        raise HTTPException(403, "Forbidden (not your file)")
    
    filename = rec.get("filename") or "file.pdf"
    safe_name = "".join(ch if ord(ch) < 128 else "_" for ch in filename)
    if not safe_name:
        safe_name = "file.pdf"

    return Response(
        content=rec["file_data"],
        media_type=rec.get("mime") or "application/pdf",
        headers={
            "Content-Disposition": f'inline; filename="{safe_name}"'
        },
    )

@files_router.delete("/declarations/{decl_id}/files/{file_id}")
def api_unlink_file(decl_id: int, file_id: int):
    deleted = unlink_file_from_declaration(decl_id, file_id)
    if not deleted:
        raise HTTPException(404, "Связь декларации и файла не найдена")
    return {"status": "ok", "deleted": deleted}

@users_router.post("/{user_id}/avatar", response_model=AvatarUploadResp)
async def upload_avatar(user_id: int, file: UploadFile = File(...), current=Depends(get_current_user)):
    require_owner(user_id, current)
    user = get_user_by_id(user_id)
    if not user:
        raise HTTPException(status_code=404, detail="User not found")

    content_type = (file.content_type or "").lower()
    if not content_type.startswith("image/"):
        raise HTTPException(status_code=400, detail="Ожидается файл изображения")

    data = await file.read()
    if not data:
        raise HTTPException(status_code=400, detail="Пустой файл")
    max_size = 5 * 1024 * 1024
    if len(data) > max_size:
        raise HTTPException(status_code=400, detail="Слишком большой файл аватара (>5 МБ)")
    mime = content_type or "application/octet-stream"
    file_id = add_file(user_id, file.filename, mime, data)

    avatar_path = f"/files/{file_id}/download"
    update_user(user_id, avatar_path=avatar_path)

    return AvatarUploadResp(file_id=file_id, avatar_path=avatar_path)




class TnvedRequestOut(BaseModel):
    id: int
    user_id: int
    query_text: Optional[str] = None
    input_json: Optional[dict] = None
    result_json: Optional[dict] = None
    status: str
    error_text: Optional[str] = None
    credits_spent: int
    model: Optional[str] = None
    latency_ms: Optional[int] = None
    created_at: Optional[str] = None  # можно и datetime, но строкой проще на фронте


class PaymentHistoryOut(BaseModel):
    id: int
    user_id: int
    provider: str
    provider_payment_id: Optional[str] = None
    tariff_plan_id: Optional[int] = None
    amount_value: str
    currency: str
    status: str
    paid_at: Optional[str] = None
    created_at: Optional[str] = None

    tariff_code: Optional[str] = None
    tariff_title: Optional[str] = None
    tariff_credits: Optional[int] = None
    tariff_price_rub: Optional[int] = None


@users_router.get("/{user_id}/tnved-requests", response_model=list[TnvedRequestOut])
def api_list_tnved_requests(user_id: int, limit: int = 200, offset: int = 0, current=Depends(get_current_user)):
    require_owner(user_id, current)
    rows = tnved_requests_list_by_user(user_id, limit=limit, offset=offset) or []
    out = []
    for r in rows:
        rr = dict(r)
        if rr.get("created_at") is not None:
            rr["created_at"] = rr["created_at"].isoformat()
        out.append(rr)
    return out


@users_router.get("/{user_id}/payments", response_model=list[PaymentHistoryOut])
def api_list_payments(user_id: int, limit: int = 200, offset: int = 0, current=Depends(get_current_user)):
    require_owner(user_id, current)
    rows = payments_list_by_user(user_id, limit=limit, offset=offset) or []
    out = []
    for r in rows:
        rr = dict(r)
        if rr.get("created_at") is not None:
            rr["created_at"] = rr["created_at"].isoformat()
        if rr.get("paid_at") is not None:
            rr["paid_at"] = rr["paid_at"].isoformat()
        out.append(rr)
    return out


######################## Админ-панель #####################################

class AdminUserRow(BaseModel):
    id: int
    email: Optional[str] = None
    name: Optional[str] = None
    surname: Optional[str] = None
    avatar_path: Optional[str] = None


class AdminDeclarationRow(BaseModel):
    id: int
    user_id: int
    title: str
    created_at: Optional[datetime] = None


class AdminDeclarationDetails(BaseModel):
    id: int
    user_id: int
    title: str
    created_at: Optional[datetime] = None


class AdminJobRow(BaseModel):
    id: int
    decl_id: int
    file_id: int
    doc_key: Optional[str] = None
    status: str
    attempts: Optional[int] = None
    error_text: Optional[str] = None
    created_at: Optional[datetime] = None


class AdminJobDetails(AdminJobRow):
    pass


def as_dict(row: Any, keys: List[str]) -> Dict[str, Any]:
    """
    Поддержка и dict-строк (RealDictCursor), и tuple-строк.
    """
    if row is None:
        return {}
    if isinstance(row, dict):
        return row
    # tuple/list
    return {k: row[i] for i, k in enumerate(keys) if i < len(row)}


@admin_router.get("/health")
def admin_health():
    return {"status": "ok"}


@admin_router.get("/users", response_model=List[AdminUserRow])
def admin_users(q: str = "", limit: int = 50, offset: int = 0):
    q_like = f"%{q}%"

    with get_conn() as conn:
        cur = conn.cursor()
        cur.execute(
            """
            SELECT id, email, name, surname, avatar_path
            FROM users
            WHERE (%s = '' OR email ILIKE %s OR name ILIKE %s OR surname ILIKE %s)
            ORDER BY id DESC
            LIMIT %s OFFSET %s
            """,
            (q, q_like, q_like, q_like, limit, offset),
        )
        rows = cur.fetchall()
        cur.close()

    out: List[AdminUserRow] = []
    keys = ["id", "email", "name", "surname", "avatar_path"]
    for r in rows:
        d = as_dict(r, keys)
        out.append(
            AdminUserRow(
                id=int(d["id"]),
                email=d.get("email"),
                name=d.get("name"),
                surname=d.get("surname"),
                avatar_path=(d.get("avatar_path") or ""),
            )
        )
    return out


@admin_router.get("/users/{user_id}", response_model=AdminUserRow)
def admin_user_details(user_id: int):
    with get_conn() as conn:
        cur = conn.cursor()
        cur.execute(
            """
            SELECT id, email, name, surname, avatar_path
            FROM users
            WHERE id = %s
            """,
            (user_id,),
        )
        r = cur.fetchone()
        cur.close()

    if not r:
        raise HTTPException(404, "User not found")

    d = as_dict(r, ["id", "email", "name", "surname", "avatar_path"])
    return AdminUserRow(
        id=int(d["id"]),
        email=d.get("email"),
        name=d.get("name"),
        surname=d.get("surname"),
        avatar_path=(d.get("avatar_path") or ""),
    )


@admin_router.get("/declarations", response_model=List[AdminDeclarationRow])
def admin_declarations(user_id: int = 0, q: str = "", limit: int = 50, offset: int = 0):
    q_like = f"%{q}%"

    with get_conn() as conn:
        cur = conn.cursor()
        if user_id:
            cur.execute(
                """
                SELECT id, user_id, title, created_at
                FROM declarations
                WHERE user_id = %s AND (%s = '' OR title ILIKE %s)
                ORDER BY id DESC
                LIMIT %s OFFSET %s
                """,
                (user_id, q, q_like, limit, offset),
            )
        else:
            cur.execute(
                """
                SELECT id, user_id, title, created_at
                FROM declarations
                WHERE (%s = '' OR title ILIKE %s)
                ORDER BY id DESC
                LIMIT %s OFFSET %s
                """,
                (q, q_like, limit, offset),
            )

        rows = cur.fetchall()
        cur.close()

    out: List[AdminDeclarationRow] = []
    keys = ["id", "user_id", "title", "created_at"]
    for r in rows:
        d = as_dict(r, keys)
        out.append(
            AdminDeclarationRow(
                id=int(d["id"]),
                user_id=int(d["user_id"]),
                title=d.get("title") or "",
                created_at=d.get("created_at"),
            )
        )
    return out


@admin_router.get("/declarations/{decl_id}", response_model=AdminDeclarationDetails)
def admin_declaration_details(decl_id: int):
    with get_conn() as conn:
        cur = conn.cursor()
        cur.execute(
            """
            SELECT id, user_id, title, created_at
            FROM declarations
            WHERE id = %s
            """,
            (decl_id,),
        )
        r = cur.fetchone()
        cur.close()

    if not r:
        raise HTTPException(404, "Declaration not found")

    d = as_dict(r, ["id", "user_id", "title", "created_at"])
    return AdminDeclarationDetails(
        id=int(d["id"]),
        user_id=int(d["user_id"]),
        title=d.get("title") or "",
        created_at=d.get("created_at"),
    )


@admin_router.get("/jobs", response_model=List[AdminJobRow])
def admin_jobs(status: str = "", decl_id: int = 0, limit: int = 50, offset: int = 0):
    where = []
    params: List[Any] = []

    if status:
        where.append("status = %s")
        params.append(status)

    if decl_id:
        where.append("decl_id = %s")
        params.append(decl_id)

    where_sql = ""
    if where:
        where_sql = "WHERE " + " AND ".join(where)

    params.extend([limit, offset])

    with get_conn() as conn:
        cur = conn.cursor()
        cur.execute(
            f"""
            SELECT id, decl_id, file_id, doc_key, status, attempts, error_text, created_at
            FROM jobs
            {where_sql}
            ORDER BY id DESC
            LIMIT %s OFFSET %s
            """,
            tuple(params),
        )
        rows = cur.fetchall()
        cur.close()

    out: List[AdminJobRow] = []
    keys = ["id", "decl_id", "file_id", "doc_key", "status", "attempts", "error_text", "created_at"]
    for r in rows:
        d = as_dict(r, keys)
        out.append(
            AdminJobRow(
                id=int(d["id"]),
                decl_id=int(d["decl_id"]),
                file_id=int(d["file_id"]),
                doc_key=d.get("doc_key"),
                status=d.get("status") or "",
                attempts=d.get("attempts"),
                error_text=d.get("error_text"),
                created_at=d.get("created_at"),
            )
        )
    return out


@admin_router.get("/jobs/{job_id}", response_model=AdminJobDetails)
def admin_job_details(job_id: int):
    with get_conn() as conn:
        cur = conn.cursor()
        cur.execute(
            """
            SELECT id, decl_id, file_id, doc_key, status, attempts, error_text, created_at
            FROM jobs
            WHERE id = %s
            """,
            (job_id,),
        )
        r = cur.fetchone()
        cur.close()

    if not r:
        raise HTTPException(404, "Job not found")

    d = as_dict(r, ["id", "decl_id", "file_id", "doc_key", "status", "attempts", "error_text", "created_at"])
    return AdminJobDetails(
        id=int(d["id"]),
        decl_id=int(d["decl_id"]),
        file_id=int(d["file_id"]),
        doc_key=d.get("doc_key"),
        status=d.get("status") or "",
        attempts=d.get("attempts"),
        error_text=d.get("error_text"),
        created_at=d.get("created_at"),
    )

@admin_router.get("/jobs/{job_id}/result")
def admin_job_result(job_id: int):
    row = jobs_get(job_id)
    if not row:
        raise HTTPException(404, "Job not found")

    status = row.get("status") or ""
    decl_id = row.get("decl_id")
    doc_key = row.get("doc_key") or ""
    result = row.get("result_json")

    if (result is None) and (status == "done") and decl_id and doc_key:
        dk = f"{doc_key}_json"
        linked = list_declaration_files(int(decl_id)) or []
        file_id = None
        for r in linked:
            if str(r.get("doc_key") or "") == dk:
                file_id = r.get("file_id")
                break

        if file_id:
            rec = get_file(int(file_id))
            if rec and rec.get("file_data"):
                try:
                    raw = rec["file_data"]
                    if isinstance(raw, (bytes, bytearray)):
                        raw = raw.decode("utf-8", errors="replace")
                    result = json.loads(raw)
                except Exception:
                    # если файл невалидный JSON — вернём как текст
                    result = {"_error": "invalid_json_file", "raw": raw}

    return {
        "job_id": row.get("id"),
        "decl_id": decl_id,
        "doc_key": doc_key,
        "status": status,
        "result": result,
    }


######################## ЮKASSA ########################

import uuid
from decimal import Decimal, ROUND_HALF_UP
from typing import Any, Dict, Optional
from fastapi import Request

YOOKASSA_BASE = "https://api.yookassa.ru/v3"
YOOKASSA_SHOP_ID = os.getenv("YOOKASSA_SHOP_ID", "")
YOOKASSA_SECRET_KEY = os.getenv("YOOKASSA_SECRET_KEY", "")
YOOKASSA_RETURN_URL = os.getenv("YOOKASSA_RETURN_URL", "")  
YOOKASSA_WEBHOOK_TOKEN = os.getenv("YOOKASSA_WEBHOOK_TOKEN", "")

if APP_ENV == "prod":
    if not JWT_SECRET:
        raise RuntimeError("JWT_SECRET is required in prod")
    if not YOOKASSA_WEBHOOK_TOKEN:
        raise RuntimeError("YOOKASSA_WEBHOOK_TOKEN is required in prod")

def _require_yookassa_config():
    if not YOOKASSA_SHOP_ID or not YOOKASSA_SECRET_KEY:
        raise HTTPException(500, "YOOKASSA_SHOP_ID/YOOKASSA_SECRET_KEY не настроены")

def _money_str(value: Decimal) -> str:
    q = value.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    return f"{q:.2f}"

def _append_query(url: str, key: str, value: str) -> str:
    if "?" in url:
        return f"{url}&{key}={value}"
    return f"{url}?{key}={value}"

async def yk_request(
    method: str,
    path: str,
    *,
    json_body: Optional[Dict[str, Any]] = None,
    idempotence_key: Optional[str] = None,
) -> Dict[str, Any]:
    _require_yookassa_config()
    headers = {"Content-Type": "application/json"}
    if idempotence_key:
        headers["Idempotence-Key"] = idempotence_key

    try:
        async with httpx.AsyncClient(timeout=30.0) as client:
            r = await client.request(
                method=method,
                url=f"{YOOKASSA_BASE}{path}",
                auth=(YOOKASSA_SHOP_ID, YOOKASSA_SECRET_KEY), 
                headers=headers,
                json=json_body,
            )
    except Exception as e:
        raise HTTPException(502, f"YooKassa недоступна: {e}")

    if r.status_code >= 400:
        try:
            err = r.json()
        except Exception:
            err = {"error": r.text}
        raise HTTPException(r.status_code, {"yookassa_error": err})

    return r.json()



@app.get("/v1/tariffs")
def api_list_tariffs():
    return {"items": list_active_tariff_plans()}


class CreatePaymentSimpleIn(BaseModel):
    user_id: int
    tariff_code: str

class CreatePaymentSimpleOut(BaseModel):
    order_id: int            
    status: str
    confirmation_url: Optional[str] = None
    provider_payment_id: Optional[str] = None

@app.post("/v1/payments/create", response_model=CreatePaymentSimpleOut)
async def create_payment_simple(body: CreatePaymentSimpleIn):
    tariff = get_tariff_plan_by_code(body.tariff_code)
    if not tariff:
        raise HTTPException(400, f"Неизвестный тариф: {body.tariff_code}")

    if not YOOKASSA_RETURN_URL:
        raise HTTPException(500, "YOOKASSA_RETURN_URL не задан. Укажи URL фронта /return")

    amount_value = _money_str(Decimal(tariff["price_rub"]))
    currency = "RUB"

    idem = str(uuid.uuid4())

    p = payments_create_pending(
        user_id=body.user_id,
        tariff_plan_id=int(tariff["id"]),
        amount_value=amount_value,
        currency=currency,
        idempotence_key=idem,
        raw_json={"stage": "created_in_db", "tariff_code": body.tariff_code},
    )
    order_id = int(p["id"])
    return_url = _append_query(YOOKASSA_RETURN_URL, "order_id", str(order_id))

    payload: Dict[str, Any] = {
        "amount": {"value": amount_value, "currency": currency},
        "confirmation": {"type": "redirect", "return_url": return_url},
        "capture": True,
        "description": f"Покупка тарифа {tariff['title']} (order #{order_id})",
        "metadata": {
            "order_id": str(order_id),
            "user_id": str(body.user_id),
            "tariff_code": str(body.tariff_code),
        },
    }

    yk = await yk_request("POST", "/payments", json_body=payload, idempotence_key=idem)

    provider_payment_id = yk.get("id")
    conf = (yk.get("confirmation") or {})
    confirmation_url = conf.get("confirmation_url")

    if not provider_payment_id or not confirmation_url:
        raise HTTPException(502, "YooKassa не вернула payment_id/confirmation_url")
    
    payments_set_provider_payment_id(order_id, provider_payment_id, yk)

    return CreatePaymentSimpleOut(
        order_id=order_id,
        status=yk.get("status") or "pending",
        confirmation_url=confirmation_url,
        provider_payment_id=provider_payment_id,
    )

class PaymentStatusOut(BaseModel):
    order_id: int
    status: str
    credits_added: int
    price_rub: int
    tariff_code: str
    credits_balance: int

@app.get("/v1/payments/orders/{order_id}", response_model=PaymentStatusOut)
async def get_order_status(order_id: int, user_id: int = Query(...)):
    p = payments_get_by_id_and_user(order_id, user_id)
    if not p:
        raise HTTPException(404, "Заказ не найден")

    provider_pid = p.get("provider_payment_id")
    status = p.get("status") or "pending"

    if provider_pid and status in ("pending", "waiting_for_capture"):
        yk = await yk_request("GET", f"/payments/{provider_pid}")
        new_status = yk.get("status") or status
        if new_status != status:
            paid_at = datetime.now(APP_TZ) if new_status == "succeeded" else None
            updated = payments_update_status_by_provider_id(provider_pid, new_status, yk, paid_at=paid_at)
            if updated:
                p = updated
                status = new_status

                if status == "succeeded":
                    tariff = get_tariff_plan_by_code((yk.get("metadata") or {}).get("tariff_code", "") or "")
                    if tariff:
                        credits_apply_purchase(
                            user_id=user_id,
                            payment_id=order_id,
                            credits=int(tariff["credits"]),
                            meta={"source": "status_sync"},
                        )

    tariffs = list_active_tariff_plans()
    t = next((x for x in tariffs if int(x["id"]) == int(p.get("tariff_plan_id") or 0)), None)
    if not t:
        t = {"code": "unknown", "price_rub": int(Decimal(p["amount_value"])), "credits": 0}

    balance = credits_get_balance(user_id)

    return PaymentStatusOut(
        order_id=order_id,
        status=status,
        credits_added=int(t["credits"]),
        price_rub=int(t["price_rub"]),
        tariff_code=str(t["code"]),
        credits_balance=int(balance),
    )


@app.post("/v1/yookassa/webhook")
async def yookassa_webhook(request: Request):
    token_q = request.query_params.get("token")
    token_h = request.headers.get("X-Webhook-Token")

    if YOOKASSA_WEBHOOK_TOKEN:
        if token_q != YOOKASSA_WEBHOOK_TOKEN and token_h != YOOKASSA_WEBHOOK_TOKEN:
            raise HTTPException(401, "Invalid webhook token")

    payload = await request.json()

    event = payload.get("event") 
    obj = payload.get("object") or {}
    provider_payment_id = obj.get("id")
    status = obj.get("status") or ""

    if not provider_payment_id:
        raise HTTPException(400, "No payment id in webhook")

    yk_actual = await yk_request("GET", f"/payments/{provider_payment_id}")
    actual_status = yk_actual.get("status") or status

    paid_at = datetime.now(APP_TZ) if actual_status == "succeeded" else None
    updated = payments_update_status_by_provider_id(provider_payment_id, actual_status, yk_actual, paid_at=paid_at)

    if not updated:
        return {"ok": True, "ignored": True}

    if actual_status == "succeeded":
        order_id = int(updated["id"])
        user_id = int(updated["user_id"])

        tariffs = list_active_tariff_plans()
        t = next((x for x in tariffs if int(x["id"]) == int(updated.get("tariff_plan_id") or 0)), None)
        if t:
            credits_apply_purchase(
                user_id=user_id,
                payment_id=order_id,
                credits=int(t["credits"]),
                meta={"source": "webhook", "event": event},
            )

    return {"ok": True, "event": event, "provider_payment_id": provider_payment_id, "status": actual_status}


class CreditsStatusOut(BaseModel):
    credits_remaining: int
    days_left: int

@app.get("/v1/credits/status", response_model=CreditsStatusOut)
def credits_status(user_id: int = Query(...)):
    u = get_user_by_id(user_id)
    if not u:
        raise HTTPException(404, "User not found")

    s = credits_get_status(int(user_id))
    return CreditsStatusOut(
        credits_remaining=int(s.get("credits_remaining") or 0),
        days_left=int(s.get("days_left") or 0),
    )




######################## Сбор ALL_DATA и GRAPHS ############################

def fill_missing(a: Any, b: Any) -> Any:
    def is_scalar(v: Any) -> bool:
        return isinstance(v, (str, int, float, bool))
    def is_empty_scalar(v: Any) -> bool:
        if v is None: return True
        if isinstance(v, str): return v.strip() in ("", "null", "-", "—")
        return False
    if a is None:
        return b

    if isinstance(a, dict) and isinstance(b, dict):
        res = dict(a)
        for k, vb in b.items():
            if k not in res:
                res[k] = vb
                continue
            va = res[k]
            if is_scalar(va) and is_empty_scalar(va):
                if is_scalar(vb) and not is_empty_scalar(vb):
                    res[k] = vb
                continue
            res[k] = fill_missing(va, vb)
        return res

    if isinstance(a, list) and isinstance(b, list):
        out = []
        n = max(len(a), len(b))
        for i in range(n):
            if i < len(a) and i < len(b):
                out.append(fill_missing(a[i], b[i]))
            elif i < len(a):
                out.append(a[i])
            else:
                out.append(b[i])
        return out

    if is_scalar(a) and is_scalar(b):
        return b if is_empty_scalar(a) and not is_empty_scalar(b) else a

    if a in (None, "", [], {}):
        return b
    return a

def _normalize_date_for_cbrf(value) -> str:
    if value is None:
        return ""
    s = str(value).strip()
    if not s:
        return ""

    if re.fullmatch(r"\d{2}\.\d{2}\.\d{4}", s):
        return s

    m = re.match(r"^(\d{4})-(\d{2})-(\d{2})", s)
    if m:
        y, mm, dd = m.group(1), m.group(2), m.group(3)
        return f"{dd}.{mm}.{y}"

    return ""

def _is_empty_override(v) -> bool:
    if v is None:
        return True
    if isinstance(v, str) and v.strip().lower() in ("", "null", "-", "—"):
        return True
    return False



def build_all_data_for_decl(decl_id: int) -> Dict[str, Any]:
    linked_all = list_declaration_files(decl_id) or []

    def _inject_decl_date(all_data: Dict[str, Any]) -> Dict[str, Any]:
        decl_date = ""
        try:
            decl_date = get_declaration_date(decl_id) or ""
        except Exception:
            decl_date = ""

        if decl_date and isinstance(all_data, dict):
            all_data.setdefault("declaration", {})
            if isinstance(all_data.get("declaration"), dict):
                all_data["declaration"]["Дата декларации"] = decl_date
                all_data["declaration"]["date"] = decl_date
        return all_data

    all_data_link = next((r for r in linked_all if str(r.get("doc_key") or "") == "all_data_json"), None)

    if all_data_link:
        rec = get_file(all_data_link["file_id"])
        if not rec or not rec.get("file_data"):
            return {}
        try:
            all_data = json.loads(rec["file_data"].decode("utf-8", errors="ignore"))
            if isinstance(all_data, dict):
                return _inject_decl_date(all_data)
            return {}
        except Exception:
            return {}

    groups: Dict[str, list[Dict[str, Any]]] = {}
    for r in linked_all:
        dk = str(r.get("doc_key") or "")
        if dk.endswith("_json"):
            base = dk[:-5]  # invoice_json -> invoice
            groups.setdefault(base, []).append(r)

    combined: Dict[str, Any] = {}

    for base, rows in groups.items():
        rows_sorted = sorted(rows, key=lambda x: str(x.get("created_at") or ""))
        acc = None
        for row in rows_sorted:
            rec = get_file(row["file_id"])
            if not rec or not rec.get("file_data"):
                continue
            try:
                obj = json.loads(rec["file_data"].decode("utf-8", errors="ignore"))
            except Exception:
                continue
            if not isinstance(obj, (dict, list)):
                continue
            acc = obj if acc is None else fill_missing(acc, obj)
        combined[base] = acc or {}

    all_data = {k: v for k, v in combined.items() if v}
    return _inject_decl_date(all_data)

def _select_primary_tnved(all_data: Dict[str, Any]) -> str:
    from graph import get_tnved

    codes_raw = get_tnved(all_data)
    if isinstance(codes_raw, set):
        codes = sorted(c for c in codes_raw if c)
        return codes[0] if codes else ""

    if isinstance(codes_raw, (list, tuple)):
        for c in codes_raw:
            c = (c or "").strip()
            if c:
                return c
        return ""

    if isinstance(codes_raw, str):
        return codes_raw.strip()

    return ""

def _collect_tnved_list(all_data: Dict[str, Any]) -> List[str]:
    from graph import get_tnved

    codes_raw = get_tnved(all_data)
    result: List[str] = []

    if isinstance(codes_raw, set):
        result = sorted(c for c in codes_raw if c)
    elif isinstance(codes_raw, (list, tuple)):
        for c in codes_raw:
            c = (c or "").strip()
            if c:
                result.append(c)
    elif isinstance(codes_raw, str):
        c = codes_raw.strip()
        if c:
            result = [c]
    seen = set()
    uniq: List[str] = []
    for c in result:
        if c not in seen:
            seen.add(c)
            uniq.append(c)

    return uniq

def compute_date_declararion(all_data: Dict[str, Any], overrides: Dict[str, Any]) -> Dict[str, Any]:
    declaration_date = overrides.get("declaration_date") or all_data.get("declaration", {}).get("Дата декларации")
    return {"declaration_date": declaration_date}

def compute_g1(all_data: Dict[str, Any], overrides: Dict[str, Any]) -> Dict[str, Any]:
    from graph import normalize_country, get_any
    poluchatel_country = normalize_country(get_any(all_data, [
        "invoice.Получатель.Страна",
        "contract.Общая информация.Стороны.Получатель.Юридический адрес.Страна",
    ]))
    otpravitel_country = normalize_country(get_any(all_data, [
        "invoice.Отправитель.Страна",
        "contract.Общая информация.Стороны.Отправитель.Страна",
    ]))

    default_kind_tp = ""
    default_code_tp = ""
    if poluchatel_country and poluchatel_country != otpravitel_country:
        default_kind_tp = "ИМ"
        default_code_tp = "40"

    g1_1 = overrides.get("g1_1", default_kind_tp)
    g1_2 = overrides.get("g1_2", default_code_tp)
    g1_3 = overrides.get("g1_3", "ЭД")

    return {
        "g1_1": g1_1,
        "g1_2": g1_2,
        "g1_3": g1_3,
    }

def compute_g2(all_data: Dict[str, Any], overrides: Dict[str, Any]) -> Dict[str, Any]:
    from graph import (
        get_any,
        normalize_country,
        get_country_code,
        extract_index,
    )
    sender_country_norm = normalize_country(get_any(all_data, [
        "invoice.Отправитель.Страна",
        "contract.Общая информация.Стороны.Отправитель.Страна",
    ]))
    default_g2_1 = get_any(all_data, [
        "invoice.Отправитель.ИНН",
        "contract.Общая информация.Стороны.Отправитель.ИНН",
    ])
    default_g2_2 = get_any(all_data, [
        "invoice.Отправитель.КПП",
        "contract.Общая информация.Стороны.Отправитель.КПП",
    ])
    default_g2_3 = get_any(all_data, [
        "invoice.Отправитель.Название компании",
        "contract.Общая информация.Стороны.Отправитель.Название компании",
    ])
    default_g2_addr_invoice = get_any(all_data, [
        "invoice.Отправитель.Юридический адрес.Полностью",
    ])

    default_g2_addr_contract = get_any(all_data, [
        "contract.Общая информация.Стороны.Отправитель.Юридический адрес.Полностью",
    ])
    default_g2_5 = sender_country_norm or get_any(all_data, [
        "invoice.Отправитель.Страна",
        "contract.Общая информация.Стороны.Отправитель.Страна",
    ])
    default_g2_6 = extract_index(get_any(all_data, [
        "invoice.Отправитель.Юридический адрес.Полностью",
        "contract.Общая информация.Стороны.Отправитель.Юридический адрес.Полностью",
    ]))
    default_g2_7 = get_any(all_data, [
        "invoice.Отправитель.Юридический адрес.Регион/Область",
        "contract.Общая информация.Стороны.Отправитель.Юридический адрес.Регион/Область",
    ])
    default_g2_8 = get_any(all_data, [
        "invoice.Отправитель.Юридический адрес.Город",
        "contract.Общая информация.Стороны.Отправитель.Юридический адрес.Город",
    ])
    default_g2_9 = get_any(all_data, [
        "invoice.Отправитель.Юридический адрес.Улица",
        "contract.Общая информация.Стороны.Отправитель.Юридический адрес.Улица",
    ])
    default_g2_10 = get_any(all_data, [
        "invoice.Отправитель.Юридический адрес.Номер дома",
        "contract.Общая информация.Стороны.Отправитель.Юридический адрес.Номер дома",
    ])
    default_g2_11 = get_any(all_data, [
        "contract.Общая информация.Стороны.Отправитель.ОГРН",
    ])
    g2_1 = overrides.get("g2_1", default_g2_1)
    g2_2 = overrides.get("g2_2", default_g2_2)
    g2_3 = overrides.get("g2_3", default_g2_3)
    g2_5 = overrides.get("g2_5", default_g2_5)
    if "g2_4" in overrides:
        g2_4 = overrides["g2_4"]
    else:
        g2_4 = get_country_code(g2_5) if g2_5 else ""

    g2_6 = overrides.get("g2_6", default_g2_6)
    g2_7 = overrides.get("g2_7", default_g2_7)
    g2_8 = overrides.get("g2_8", default_g2_8)
    g2_9 = overrides.get("g2_9", default_g2_9)
    g2_10 = overrides.get("g2_10", default_g2_10)
    g2_11 = overrides.get("g2_11", default_g2_11)

    g2_addr_invoice = overrides.get("g2_addr_invoice", default_g2_addr_invoice)
    g2_addr_contract = overrides.get("g2_addr_contract", default_g2_addr_contract)

    return {
        "g2_1": g2_1,
        "g2_2": g2_2,
        "g2_3": g2_3,
        "g2_4": g2_4,
        "g2_5": g2_5,
        "g2_6": g2_6,
        "g2_7": g2_7,
        "g2_8": g2_8,
        "g2_9": g2_9,
        "g2_10": g2_10,
        "g2_11": g2_11,
        "g2_addr_invoice": g2_addr_invoice,
        "g2_addr_contract": g2_addr_contract
    }

def compute_g3(all_data: Dict[str, Any], overrides: Dict[str, Any]) -> Dict[str, Any]:
    from graph import get_tnved
    default_g3_1 = "1"
    codes_raw = get_tnved(all_data)
    if isinstance(codes_raw, set):
        pages = len([c for c in codes_raw if c])
    else:
        pages = 0

    default_g3_2 = str(pages)
    g3_1 = overrides.get("g3_1", default_g3_1)
    g3_2 = overrides.get("g3_2", default_g3_2)

    return {
        "g3_1": g3_1,
        "g3_2": g3_2,
    }

def compute_g4(all_data: Dict[str, Any], overrides: Dict[str, Any]) -> Dict[str, Any]:
    default_g4_1 = ""
    default_g4_2 = ""

    g4_1 = overrides.get("g4_1", default_g4_1)
    g4_2 = overrides.get("g4_2", default_g4_2)

    return {
        "g4_1": g4_1,
        "g4_2": g4_2,
    }

def compute_g5(all_data: Dict[str, Any], overrides: Dict[str, Any]) -> Dict[str, Any]:
    from graph import get_tnved 
    default_g5_1 = ""
    try:
        codes_raw = get_tnved(all_data)
        if isinstance(codes_raw, set):
            count = len([c for c in codes_raw if c])
        elif codes_raw:
            try:
                count = len(codes_raw)
            except Exception:
                count = 0
        else:
            count = 0
        default_g5_1 = str(count)
    except Exception:
        default_g5_1 = ""

    g5_1 = overrides.get("g5_1", default_g5_1)

    return {
        "g5_1": g5_1,
    }

def compute_g6(all_data: Dict[str, Any], overrides: Dict[str, Any]) -> Dict[str, Any]:
    from graph import get_total_places

    try:
        places_raw = get_total_places(all_data)
    except Exception:
        places_raw = ""
    if places_raw is None:
        default_g6_1 = ""
    else:
        default_g6_1 = str(places_raw)

    g6_1 = overrides.get("g6_1", default_g6_1)

    return {
        "g6_1": g6_1,
    }

def compute_g7(all_data: Dict[str, Any], overrides: Dict[str, Any]) -> Dict[str, Any]:
    default_g7_1 = ""
    g7_1 = overrides.get("g7_1", default_g7_1)

    return {
        "g7_1": g7_1,
    }

def compute_g8(all_data: Dict[str, Any], overrides: Dict[str, Any]) -> Dict[str, Any]:
    from graph import (
        get_any,
        normalize_country,
        get_country_code,
        extract_index,
    )

    default_g8_1 = ""
    default_g8_2 = ""
    default_g8_3 = "СМ. ГРАФУ 14 ДТ"
    default_g8_4 = ""
    default_g8_5 = ""
    default_g8_6 = ""
    default_g8_7 = ""
    default_g8_8 = ""
    default_g8_9 = ""
    default_g8_10 = ""
    default_g8_11 = ""

    g8_1 = overrides.get("g8_1", default_g8_1)
    g8_2 = overrides.get("g8_2", default_g8_2)
    g8_3 = overrides.get("g8_3", default_g8_3)
    g8_5 = overrides.get("g8_5", default_g8_5)

    if "g8_4" in overrides:
        g8_4 = overrides["g8_4"]
    else:
        g8_4 = get_country_code(g8_5) if g8_5 else default_g8_4

    g8_6 = overrides.get("g8_6", default_g8_6)
    g8_7 = overrides.get("g8_7", default_g8_7)
    g8_8 = overrides.get("g8_8", default_g8_8)
    g8_9 = overrides.get("g8_9", default_g8_9)
    g8_10 = overrides.get("g8_10", default_g8_10)
    g8_11 = overrides.get("g8_11", default_g8_11)

    return {
        "g8_1": g8_1,
        "g8_2": g8_2,
        "g8_3": g8_3,
        "g8_4": g8_4,
        "g8_5": g8_5,
        "g8_6": g8_6,
        "g8_7": g8_7,
        "g8_8": g8_8,
        "g8_9": g8_9,
        "g8_10": g8_10,
        "g8_11": g8_11,
    }

def compute_g9(all_data: Dict[str, Any], overrides: Dict[str, Any]) -> Dict[str, Any]:
    from graph import (
        get_any,
        normalize_country,
        get_country_code,
        extract_index,
    )


    default_g9_1 = ""
    default_g9_2 = ""
    default_g9_3 = "СМ. ГРАФУ 14 ДТ"
    default_g9_4 = ""
    default_g9_5 = ""
    default_g9_6 = ""
    default_g9_7 = ""
    default_g9_8 = ""
    default_g9_9 = ""
    default_g9_10 = ""
    default_g9_11 = ""

    g9_1 = overrides.get("g9_1", default_g9_1)
    g9_2 = overrides.get("g9_2", default_g9_2)
    g9_3 = overrides.get("g9_3", default_g9_3)
    g9_5 = overrides.get("g9_5", default_g9_5)
    if "g9_4" in overrides:
        g9_4 = overrides["g9_4"]
    else:
        g9_4 = get_country_code(g9_5) if g9_5 else default_g9_4

    g9_6 = overrides.get("g9_6", default_g9_6)
    g9_7 = overrides.get("g9_7", default_g9_7)
    g9_8 = overrides.get("g9_8", default_g9_8)
    g9_9 = overrides.get("g9_9", default_g9_9)
    g9_10 = overrides.get("g9_10", default_g9_10)
    g9_11 = overrides.get("g9_11", default_g9_11)

    return {
        "g9_1": g9_1,
        "g9_2": g9_2,
        "g9_3": g9_3,
        "g9_4": g9_4,
        "g9_5": g9_5,
        "g9_6": g9_6,
        "g9_7": g9_7,
        "g9_8": g9_8,
        "g9_9": g9_9,
        "g9_10": g9_10,
        "g9_11": g9_11,
    }
    
def compute_g11(all_data: Dict[str, Any], overrides: Dict[str, Any]) -> Dict[str, Any]:
    from graph import get_any, normalize_country, get_country_code
    sender_country = normalize_country(get_any(all_data, [
        "invoice.Отправитель.Страна",
        "contract.Общая информация.Стороны.Отправитель.Страна",
    ]))
    base_code = get_country_code(sender_country) if sender_country else ""
    g2_4_override = overrides.get("g2_4")
    if g2_4_override:
        base_code = g2_4_override

    default_g11_1 = base_code
    default_g11_2 = sender_country or ""

    g11_1 = overrides.get("g11_1", default_g11_1)
    return {
        "g11_1": g11_1
    }

def compute_g12(all_data: Dict[str, Any], overrides: Dict[str, Any]) -> Dict[str, Any]:
    from graph import get_currency, _to_decimal, get_any, normalize_country, get_country_code
    from parser_cbrf import cb_rate
    from decimal import Decimal

    def _d(val) -> Decimal:
        try:
            return _to_decimal(val)
        except Exception:
            return Decimal("0")
    date_declaration = get_any(all_data, [
        "declaration.Дата декларации",
        "declaration.date",
        "dt.Дата декларации",
        "dt.date",
    ])

    goods_invoice = all_data.get("invoice", {}).get("Товары", [])
    payment = all_data.get("payment", {}) or {}
    shipping_list = payment.get("Перевозка", []) or []
    insurance_list = payment.get("Страхование", []) or []
    payment_info = payment.get("Общая информация", {}) or {}
    cur_payment = payment_info.get("Валюта документа") or ""

    sum_inv = Decimal("0")
    sum_shipping = Decimal("0")
    sum_insurance = Decimal("0")

    if isinstance(goods_invoice, list):
        for g in goods_invoice:
            if not isinstance(g, dict):
                continue
            price = _d(g.get("Цена"))
            qty = _d(g.get("Количество"))
            total = _d(g.get("Стоимость"))
            if total > 0:
                item_total = total
            else:
                item_total = price * qty
            sum_inv += item_total

    if sum_inv == 0:
        return {
            "g12_currency": overrides.get("g12_currency", "RUB"),
            "g12_logistics": overrides.get("g12_logistics", ""),
            "g12_insurance": overrides.get("g12_insurance", ""),
            "g12_1": overrides.get("g12_1", ""),
        }
    if isinstance(shipping_list, list):
        for s in shipping_list:
            if not isinstance(s, dict):
                continue

            usluga = s.get("Услуга") or {}
            if not isinstance(usluga, dict):
                usluga = {}

            route = s.get("Маршрут") or {}
            if not isinstance(route, dict):
                route = {}

            descr = str(usluga.get("Описание") or "").lower()
            path  = str(route.get("Откуда") or "").lower()
            if "продолж" in descr or "границ" in path:
                continue

            amount = _d(usluga.get("Сумма"))
            if "страхов" in descr:
                sum_insurance += amount
            else:
                sum_shipping += amount

    if isinstance(insurance_list, list):
        for i in insurance_list:
            if not isinstance(i, dict):
                continue

            usluga = i.get("Услуга") or {}
            if not isinstance(usluga, dict):
                usluga = {}

            route = i.get("Маршрут") or {}
            if not isinstance(route, dict):
                route = {}

            descr = str(usluga.get("Описание") or "").lower()
            path  = str(route.get("Откуда") or "").lower()

            if "продолж" in descr or "границ" in path:
                continue

            if "страхов" in descr:
                amount = _d(usluga.get("Сумма"))
                sum_insurance += amount

    inv_currency = get_currency(all_data) or "RUB"
    raw_buyer_country = get_any(all_data, [
        "payment.Покупатель (Заказчик).Страна"
    ], default="")

    buyer_country_norm = normalize_country(raw_buyer_country)
    buyer_country_code = get_country_code(buyer_country_norm) if buyer_country_norm else ""
    is_buyer_ru = (buyer_country_code == "RU")
    if is_buyer_ru: 
        cur_payment = "RUB"
    if date_declaration and inv_currency and inv_currency != "RUB":
        rate_inv = _d(cb_rate(date_declaration, inv_currency))
    else:
        rate_inv = Decimal("1")
    
    if (not is_buyer_ru and date_declaration and cur_payment != "RUB"):
        rate_payment = _d(cb_rate(date_declaration, cur_payment))
    else:
        rate_payment = Decimal("1")

    sum_inv_rub = (sum_inv * rate_inv).quantize(Decimal("0.01"))
    sum_shipping_rub = (sum_shipping * rate_payment).quantize(Decimal("0.01"))
    sum_insurance_rub = (sum_insurance * rate_payment).quantize(Decimal("0.01"))

    total_rub = sum_inv_rub + sum_shipping_rub + sum_insurance_rub

    g12_currency = overrides.get("g12_currency", "RUB")
    g12_logistics = overrides.get(
        "g12_logistics",
        (str(sum_shipping_rub) if sum_shipping_rub != 0 else ""),
    )
    g12_insurance = overrides.get(
        "g12_insurance",
        (str(sum_insurance_rub) if sum_insurance_rub != 0 else ""),
    )
    g12_1 = overrides.get(
        "g12_1",
        (str(total_rub) if total_rub != 0 else ""),
    )

    return {
        "g12_cur_payment": cur_payment, 
        "g12_buyer_country": buyer_country_norm,
        "g12_buyer_country_code": buyer_country_code,
        "g12_currency": g12_currency,  
        "g12_logistics": g12_logistics,
        "g12_insurance": g12_insurance,
        "g12_1": g12_1,
    }

def compute_g14(all_data: Dict[str, Any], overrides: Dict[str, Any]) -> Dict[str, Any]:
    from graph import (
        get_any,
        normalize_country,
        get_country_code,
        extract_index,
    )
    declarant = (get_any(all_data, [
        "contract.Общая информация.Декларант.Название компании"
    ]) or "").strip()

    poluchatel = (get_any(all_data, [
        "contract.Общая информация.Стороны.Получатель.Название компании"
    ]) or "").strip()

    poluchatel_country = normalize_country(get_any(all_data, [
        "invoice.Получатель.Страна",
        "contract.Общая информация.Стороны.Получатель.Юридический адрес.Страна",
    ]))

    same_party = (
        declarant != "" and poluchatel != "" and poluchatel.upper() == declarant.upper()
    )

    if not same_party:
        default_g14_1 = get_any(all_data, [
            "invoice.Получатель.ИНН",
            "contract.Общая информация.Стороны.Получатель.ИНН",
        ])
        default_g14_2 = get_any(all_data, [
            "invoice.Получатель.КПП",
            "contract.Общая информация.Стороны.Получатель.КПП",
        ])
        default_g14_3 = get_any(all_data, [
            "invoice.Получатель.Название компании",
            "contract.Общая информация.Стороны.Получатель.Название компании",
        ])
        default_g14_addr_invoice = get_any(all_data, [
            "invoice.Получатель.Юридический адрес.Полностью",
            ])

        default_g14_addr_contract = get_any(all_data, [
            "contract.Общая информация.Стороны.Получатель.Юридический адрес.Полностью",
        ])
        default_g14_5 = poluchatel_country or get_any(all_data, [
            "invoice.Получатель.Страна",
            "contract.Общая информация.Стороны.Получатель.Юридический адрес.Страна",
        ])
        default_g14_4 = get_country_code(default_g14_5) if default_g14_5 else ""
        default_g14_6 = extract_index(get_any(all_data, [
            "invoice.Получатель.Юридический адрес.Полностью",
            "contract.Общая информация.Стороны.Получатель.Юридический адрес.Полностью",
        ]))
        default_g14_7 = get_any(all_data, [
            "invoice.Получатель.Юридический адрес.Регион/Область",
            "contract.Общая информация.Стороны.Получатель.Юридический адрес.Регион/Область",
        ])
        default_g14_8 = get_any(all_data, [
            "invoice.Получатель.Юридический адрес.Город",
            "contract.Общая информация.Стороны.Получатель.Юридический адрес.Город",
        ])
        default_g14_9 = get_any(all_data, [
            "invoice.Получатель.Юридический адрес.Улица",
            "contract.Общая информация.Стороны.Получатель.Юридический адрес.Улица",
        ])
        default_g14_10 = get_any(all_data, [
            "invoice.Получатель.Юридический адрес.Номер дома",
            "contract.Общая информация.Стороны.Получатель.Юридический адрес.Номер дома",
        ])
        default_g14_11 = get_any(all_data, [
            "contract.Общая информация.Стороны.Получатель.ОГРН",
        ])
    else:
        default_g14_1 = "СМ. ГРАФУ 14 ДТ"
        default_g14_2 = ""
        default_g14_3 = ""
        default_g14_4 = ""
        default_g14_5 = ""
        default_g14_6 = ""
        default_g14_7 = ""
        default_g14_8 = ""
        default_g14_9 = ""
        default_g14_10 = ""
        default_g14_11 = ""

    g14_1 = overrides.get("g14_1", default_g14_1)
    g14_2 = overrides.get("g14_2", default_g14_2)
    g14_3 = overrides.get("g14_3", default_g14_3)
    g14_5 = overrides.get("g14_5", default_g14_5)
    if "g14_4" in overrides:
        g14_4 = overrides["g14_4"]
    else:
        g14_4 = get_country_code(g14_5) if g14_5 else default_g14_4

    g14_6 = overrides.get("g14_6", default_g14_6)
    g14_7 = overrides.get("g14_7", default_g14_7)
    g14_8 = overrides.get("g14_8", default_g14_8)
    g14_9 = overrides.get("g14_9", default_g14_9)
    g14_10 = overrides.get("g14_10", default_g14_10)
    g14_11 = overrides.get("g14_11", default_g14_11)

    g14_addr_invoice = overrides.get("g14_addr_invoice", default_g14_addr_invoice)
    g14_addr_contract = overrides.get("g14_addr_contract", default_g14_addr_contract)
    return {
        "g14_1": g14_1,
        "g14_2": g14_2,
        "g14_3": g14_3,
        "g14_4": g14_4,
        "g14_5": g14_5,
        "g14_6": g14_6,
        "g14_7": g14_7,
        "g14_8": g14_8,
        "g14_9": g14_9,
        "g14_10": g14_10,
        "g14_11": g14_11,
        "g14_addr_invoice": g14_addr_invoice,
        "g14_addr_contract": g14_addr_contract,

    }

def compute_g15(all_data: Dict[str, Any], overrides: Dict[str, Any]) -> Dict[str, Any]:
    from graph import get_any, normalize_country, get_country_code

    default_g15_2 = normalize_country(get_any(all_data, [
        "transport_rail.Перевозка.Станция отправления.Страна",
        "transport_road.Перевозка.Место погрузки.Страна",
        "transport_air.Перевозка.Аэропорт отправления.Страна",
        "transport_sea.Перевозка.Отправитель.Страна"
    ])) or "НЕИЗВЕСТНА"

    default_g15_1 = get_country_code(default_g15_2) if default_g15_2 != "НЕИЗВЕСТНА" else ""
    g15_2 = overrides.get("g15_2", default_g15_2)

    if "g15_1" in overrides:
        g15_1 = overrides["g15_1"]
    else:
        g15_1 = get_country_code(g15_2) if g15_2 else default_g15_1

    return {
        "g15_1": g15_1,
        "g15_2": g15_2,
    }

def compute_g16(all_data: Dict[str, Any], overrides: Dict[str, Any]) -> Dict[str, Any]:
    from graph import (collect_origin_values,_is_unknown, _is_eu_label, get_country_code,normalize_country,)
    raw = collect_origin_values(all_data)
    norms = [(orig, (orig or "").upper().strip()) for orig in raw]

    countries = set()
    has_eu_label = False
    unknown_count = 0

    for orig, up in norms:
        if _is_unknown(orig):
            unknown_count += 1
            continue
        if _is_eu_label(orig):
            has_eu_label = True
            continue
        if up:
            countries.add(up)

    if (not countries) and (unknown_count > 0 or not raw):
        default_name = "НЕИЗВЕСТНА"
        default_code = ""
    elif (not countries) and has_eu_label:
        default_name = "ЕВРОСОЮЗ"
        default_code = ""
    elif len(countries) == 1 and not has_eu_label:
        single_upper = next(iter(countries))
        orig_display = next(
            (orig for orig, up in norms if up == single_upper),
            single_upper.title()
        )
        default_name = normalize_country(orig_display)
        try:
            default_code = get_country_code(default_name) or ""
        except Exception:
            default_code = ""
    else:
        default_name = "РАЗНЫЕ"
        default_code = ""

    g16_2 = overrides.get("g16_2", default_name)
    g16_1 = overrides.get("g16_1", (get_country_code(g16_2) or default_code) if g16_2 else default_code)

    return {
        "g16_1": g16_1,
        "g16_2": g16_2,
    }

def compute_g17(all_data: Dict[str, Any], overrides: Dict[str, Any]) -> Dict[str, Any]:
    from graph import get_any, normalize_country, get_country_code

    default_g17_2 = normalize_country(get_any(all_data, [
        "transport_rail.Перевозка.Станция назначения.Страна",
        "transport_road.Перевозка.Место разгрузки.Страна",
        "transport_air.Перевозка.Аэропорт назначения.Страна",
    ]))

    default_g17_1 = get_country_code(default_g17_2) if default_g17_2 else ""
    g17_2 = overrides.get("g17_2", default_g17_2)

    if "g17_1" in overrides:
        g17_1 = overrides["g17_1"]
    else:
        g17_1 = get_country_code(g17_2) if g17_2 else default_g17_1

    return {
        "g17_1": g17_1,
        "g17_2": g17_2,
    }

def compute_g18(all_data: Dict[str, Any], overrides: Dict[str, Any]) -> Dict[str, Any]:
    from graph import get_transport
    
    transport_type = next((k for k in all_data.keys() if k.startswith("transport_")), None)
    if transport_type not in ("transport_road", "transport_rail"):
        return {"g18_1": "", "g18_2": "", "g18_3": ""}

    from graph import get_transport
    try:
        count, nums, country = get_transport(all_data)
    except Exception:
        count, nums, country = 0, "", ""

    g18_1 = overrides.get("g18_1", str(count) if count else "")
    g18_2 = overrides.get("g18_2", nums or "")
    g18_3 = overrides.get("g18_3", country or "")

    return {
        "g18_1": g18_1, 
        "g18_2": g18_2, 
        "g18_3": g18_3
        }  

def compute_g19(all_data: Dict[str, Any], overrides: Dict[str, Any]) -> Dict[str, Any]:
    transport_type = next((k for k in all_data.keys() if k.startswith("transport_")), None)

    if transport_type in ("transport_sea", "transport_rail"):
        default = "1"
    elif transport_type in ("transport_road", "transport_air"):
        default = "0"
    else:
        default = ""

    g19_1 = overrides.get("g19_1", default)
    return {"g19_1": g19_1}

def compute_g20(all_data: Dict[str, Any], overrides: Dict[str, Any]) -> Dict[str, Any]:
    from graph import get_incoterms, get_any
    inc_code, inc_place = get_incoterms(get_any(all_data, [
                    "invoice.Общая информация.Условия поставки (Incoterms)",
                    "contract.Поставка.Условия поставки (Incoterms)"]))
    default_g20_1 = inc_code
    default_g20_2 = inc_place

    g20_1 = overrides.get("g20_1", default_g20_1)
    g20_2 = overrides.get("g20_2", default_g20_2)
    return {
        "g20_1": g20_1,
        "g20_2": g20_2}

def compute_g21(all_data: Dict[str, Any], overrides: Dict[str, Any]) -> Dict[str, Any]:
    from graph import get_transport
    def _get_transport_type(all_data: Dict[str, Any]) -> str:
        return next((k for k in (all_data or {}).keys() if str(k).startswith("transport_")), "")
    
    transport_type = _get_transport_type(all_data)
    if transport_type not in ("transport_air", "transport_sea"):
        return {"g21_1": "", "g21_2": "", "g21_3": ""}

    from graph import get_transport
    try:
        count, nums, country = get_transport(all_data)
    except Exception:
        count, nums, country = 0, "", ""

    g21_1 = overrides.get("g21_1", str(count) if count else "")
    g21_2 = overrides.get("g21_2", nums or "")
    g21_3 = overrides.get("g21_3", country or "")

    return {
        "g21_1": g21_1, 
        "g21_2": g21_2, 
        "g21_3": g21_3
        }

def compute_g22(all_data: Dict[str, Any], overrides: Dict[str, Any]) -> Dict[str, Any]:
    from graph import get_currency, get_total_sum_invoice
    default_currency = get_currency(all_data) or ""
    default_sum = get_total_sum_invoice(all_data)
    default_sum_str = str(default_sum) if default_sum is not None else ""

    g22_1 = overrides.get("g22_1", default_currency)
    g22_2 = overrides.get("g22_2", default_sum_str)

    return {
        "g22_1": g22_1,
        "g22_2": g22_2,
    }

def compute_g23(all_data: Dict[str, Any], overrides: Dict[str, Any]) -> Dict[str, Any]:
    from graph import get_any, get_currency
    from parser_cbrf import cb_rate

    raw_date = overrides.get("declaration_date")  
    if _is_empty_override(raw_date):
        raw_date = get_any(all_data, [
            "declaration.Дата декларации",
            "declaration.date",
            "dt.Дата декларации",
            "dt.date",
        ])
    date_declaration = _normalize_date_for_cbrf(raw_date)
    currency = overrides.get("g22_1")
    if _is_empty_override(currency):
        currency = get_currency(all_data)
    currency = (str(currency or "")).strip().upper()

    default_rate = ""
    if date_declaration and currency:
        try:
            r = cb_rate(date_declaration, currency)
            default_rate = "" if r in (None, "") else str(r)
        except Exception:
            default_rate = ""

    manual_rate = overrides.get("g23_1")
    g23_1 = default_rate if _is_empty_override(manual_rate) else str(manual_rate).strip()
    g23_2 = overrides.get("g23_2", currency)

    return {"g23_1": g23_1, "g23_2": g23_2}

def compute_g24(all_data: Dict[str, Any], overrides: Dict[str, Any]) -> Dict[str, Any]:
    from graph import get_any, get_currency
    from parser_cbrf import cb_rate

    default_g24_1 = "010"
    default_g24_2 = "00" 

    try:
        summa_raw = get_any(all_data, ["contract.Оплата контракта.Общая сумма"])
        date_decl = get_any(all_data, [
            "declaration.Дата декларации",
            "declaration.date",
            "dt.Дата декларации",
            "dt.date",
        ])
        rate_raw = cb_rate(date_decl, get_currency(all_data))
        summa = Decimal(str(summa_raw or "0").replace(",", "."))
        rate = Decimal(str(rate_raw or "1"))

        summa_rub = (summa * rate).quantize(Decimal("0.01"))
        default_g24_2 = "06" if summa_rub < Decimal("3000000") else "00"
    except Exception:
        pass

    return {
        "g24_1": overrides.get("g24_1", default_g24_1),
        "g24_2": overrides.get("g24_2", default_g24_2),
    }

def compute_g25(all_data: Dict[str, Any], overrides: Dict[str, Any]) -> Dict[str, Any]:
    from graph import get_transport_type, get_transport
    transport_type = next(
        (k for k in all_data.keys() if isinstance(k, str) and k.startswith("transport_")),
        None
    )
    try:
        _, nums, _ = get_transport(all_data)
    except Exception:
        nums = ""

    default = ""
    if transport_type in ("transport_sea", "transport_air", "transport_road", "transport_rail"):
        default = get_transport_type(all_data, nums) or ""

    g25_1 = overrides.get("g25_1", default)
    return {"g25_1": g25_1}

def compute_g26(all_data: Dict[str, Any], overrides: Dict[str, Any]) -> Dict[str, Any]:
    from graph import get_transport_type, get_transport

    transport_type = next(
        (k for k in all_data.keys() if isinstance(k, str) and k.startswith("transport_")),
        None
    )

    default = ""

    if transport_type in ("transport_road", "transport_rail"):
        try:
            _, nums, _ = get_transport(all_data)
        except Exception:
            nums = ""

        default = get_transport_type(all_data, nums) or ""

    g26_1 = overrides.get("g26_1", default)
    return {"g26_1": g26_1}

def compute_g29(all_data: Dict[str, Any], overrides: Dict[str, Any]) -> Dict[str, Any]:
    from graph import get_any
    tp_code = get_any(all_data, [
        "transport_rail.Таможенный пост.Код ТП",
        "transport_road.Таможенный пост.Код ТП",
        "transport_air.Таможенный пост.Код ТП",
        "transport_sea.Таможенный пост.Код ТП",
    ]) or ""


    transport_type = next((k for k in all_data.keys() if k.startswith("transport_")), None)

    if transport_type in ("transport_sea", "transport_rail"):
        default_g29_1 = str(tp_code).strip()
    else:
        default_g29_1 = ""

    default_g29_1 = str(tp_code).strip()

    g29_1 = overrides.get("g29_1", default_g29_1)

    return {
        "g29_1": g29_1
        }

def compute_g30(all_data: Dict[str, Any], overrides: Dict[str, Any]) -> Dict[str, Any]:
    from graph import get_any
    from parcer_alta_tam import get_svh_data

    # Определяем тип транспортного документа
    transport_type = next(
        (k for k in all_data.keys() if isinstance(k, str) and k.startswith("transport_")),
        None,
    )

    if transport_type == "transport_sea":
        default_g30_1 = "95"
        default_g30_2 = "2"
    elif transport_type == "transport_rail":
        default_g30_1 = "99"
        default_g30_2 = "2"
    elif transport_type in ("transport_road", "transport_air"):
        default_g30_1 = "11"
        default_g30_2 = "2"
    else:
        default_g30_1 = ""
        default_g30_2 = ""

    tp_code = get_any(all_data, [
        "transport_rail.Таможенный пост.Код ТП",
        "transport_road.Таможенный пост.Код ТП",
        "transport_air.Таможенный пост.Код ТП",
        "transport_sea.Таможенный пост.Код ТП",
    ]) or ""

    default_g30_3 = str(tp_code).strip()

    g30_1 = overrides.get("g30_1", default_g30_1)
    g30_2 = overrides.get("g30_2", default_g30_2)
    g30_3 = overrides.get("g30_3", default_g30_3)

    svh = {}
    name = ""
    license_number = ""
    license_date = ""
    address = ""
    country_code = ""
    country_name = ""
    region = ""
    city = ""
    street_house = ""

    if g30_3:
        try:
            svh = get_svh_data(g30_3) or {}
        except Exception:
            svh = {}

    if isinstance(svh, dict) and svh:
        first = next(iter(svh.values()))
        if isinstance(first, dict):
            name           = first.get("Наименование СВХ", "") or ""
            license_number = first.get("Номер лицензии", "") or ""
            license_date   = first.get("Дата лицензии", "") or ""
            address        = first.get("Адрес", "") or ""
            country_code   = first.get("CountryCode", "") or ""
            country_name   = first.get("CountryName", "") or ""
            region         = first.get("Region", "") or ""
            city           = first.get("City", "") or ""
            street_house   = first.get("StreetHouse", "") or ""

    g30_svh_name       = overrides.get("g30_svh_name", name)
    g30_license_number = overrides.get("g30_license_number", license_number)
    g30_license_date   = overrides.get("g30_license_date", license_date)
    g30_address        = overrides.get("g30_address", address)

    g30_country_code   = overrides.get("g30_country_code", country_code)
    g30_country_name   = overrides.get("g30_country_name", country_name)
    g30_region         = overrides.get("g30_region", region)
    g30_city           = overrides.get("g30_city", city)
    g30_street_house   = overrides.get("g30_street_house", street_house)

    return {
        "g30_1": g30_1,
        "g30_2": g30_2,
        "g30_3": g30_3,
        "g30_svh_name": g30_svh_name,
        "g30_license_number": g30_license_number,
        "g30_license_date": g30_license_date,
        "g30_address": g30_address,
        "g30_country_code": g30_country_code,
        "g30_country_name": g30_country_name,
        "g30_region": g30_region,
        "g30_city": g30_city,
        "g30_street_house": g30_street_house,
        "svh": svh,
    }





###################### Товарные графы ###################
def compute_g31(all_data: Dict[str, Any], overrides: Dict[str, Any]) -> Dict[str, Any]:
    from graph import get_product_country, get_seats, normalize_country

    tnved_list = _collect_tnved_list(all_data) or []
    primary_tnved = _select_primary_tnved(all_data)

    invoice = (all_data.get("invoice") or {})
    invoice_goods = invoice.get("Товары", []) or []
    if not isinstance(invoice_goods, list):
        invoice_goods = []

    def _safe_str(x: Any) -> str:
        return "" if x in (None, "") else str(x).strip()

    def _norm_code(x: Any) -> str:
        return _safe_str(x).replace(" ", "")

    def _build_desc_fallback(good: Dict[str, Any]) -> str:
        name = _safe_str(good.get("Наименование") or good.get("Описание"))
        extra_parts: List[str] = []
        for key in ("Модель", "Артикул", "Характеристики", "Маркировка"):
            v = _safe_str(good.get(key))
            if v:
                extra_parts.append(f"{key}: {v}")
        if extra_parts:
            return f"{name} — " + "; ".join(extra_parts) if name else "; ".join(extra_parts)
        return name

    by_tnved: Dict[str, Dict[str, Any]] = {}
    for g in invoice_goods:
        if not isinstance(g, dict):
            continue
        code = _norm_code(g.get("Код ТНВЭД"))
        if code and code not in by_tnved:
            by_tnved[code] = g

    groups_raw = invoice.get("_tnved_groups") or {}
    if not isinstance(groups_raw, dict):
        groups_raw = {}

    group_desc_by_code: Dict[str, str] = {}
    for k, v in groups_raw.items():
        code = _norm_code(k)
        if not code:
            continue
        desc = ""
        if isinstance(v, dict):
            desc = _safe_str(v.get("Описание группы") or v.get("group_desc") or v.get("groupDesc"))
        else:
            desc = _safe_str(v)
        if desc:
            group_desc_by_code[code] = desc

    def _desc_for_code(code_any: Any) -> str:
        code = _norm_code(code_any)
        if not code:
            return ""

        gd = group_desc_by_code.get(code)
        if gd:
            return gd
        good = by_tnved.get(code)
        if isinstance(good, dict):
            tech = _safe_str(good.get("Техническое описание"))
            if tech:
                return tech
            return _build_desc_fallback(good)
        return ""

    over_list = overrides.get("g31_1_list")
    if isinstance(over_list, (list, tuple)):
        g31_1_list = [(_safe_str(x)) for x in over_list]
        target_len = len(tnved_list) if tnved_list else len(g31_1_list)
        if target_len and len(g31_1_list) < target_len:
            g31_1_list += [""] * (target_len - len(g31_1_list))
        elif target_len and len(g31_1_list) > target_len:
            g31_1_list = g31_1_list[:target_len]
    else:
        g31_1_list: List[str] = []
        if tnved_list:
            for code in tnved_list:
                g31_1_list.append(_desc_for_code(code))
        else:
            for g in invoice_goods:
                if not isinstance(g, dict):
                    continue
                g31_1_list.append(_desc_for_code(g.get("Код ТНВЭД")) or _build_desc_fallback(g))

    if "g31_1" in overrides:
        g31_1_scalar = _safe_str(overrides.get("g31_1"))
    else:
        if primary_tnved and tnved_list:
            try:
                idx = tnved_list.index(primary_tnved)
                g31_1_scalar = g31_1_list[idx] if idx < len(g31_1_list) else ""
            except ValueError:
                g31_1_scalar = g31_1_list[0] if g31_1_list else ""
        else:
            g31_1_scalar = g31_1_list[0] if g31_1_list else ""

    good_primary = by_tnved.get(_norm_code(primary_tnved)) if primary_tnved else None
    if not isinstance(good_primary, dict) and invoice_goods:
        good_primary = invoice_goods[0] if isinstance(invoice_goods[0], dict) else None
    good_primary = good_primary or {}

    qty_main = good_primary.get("Количество")
    default_qty_1 = _safe_str(qty_main)

    default_qty_2 = ""
    default_g31_ois = ""

    try:
        seats_map = get_seats(all_data) or {}
    except Exception:
        seats_map = {}

    seats_val = seats_map.get(primary_tnved)
    default_g31_places = _safe_str(seats_val) if seats_val not in (None, "", 0) else ""

    default_g31_pallets = ""

    prod_country = get_product_country(all_data)
    if isinstance(prod_country, set):
        prod_country = next(iter(sorted(prod_country))) if prod_country else ""
    prod_country = normalize_country(prod_country) if prod_country else ""
    default_g31_origin = prod_country

    return {
        "g31_1_list": g31_1_list,
        # "g31_1": g31_1_scalar,
        # "qty_1": overrides.get("qty_1", default_qty_1),
        # "qty_2": overrides.get("qty_2", default_qty_2),
        # "g31_ois": overrides.get("g31_ois", default_g31_ois),
        # "g31_places": overrides.get("g31_places", default_g31_places),
        # "g31_pallets": overrides.get("g31_pallets", default_g31_pallets),
        # "g31_origin": overrides.get("g31_origin", default_g31_origin),
    }

def compute_g32(all_data: Dict[str, Any], overrides: Dict[str, Any]) -> Dict[str, Any]:
    default_g32_1 = "1"
    return {
        "g32_1": overrides.get("g32_1", default_g32_1),
    }

def compute_g33(all_data: Dict[str, Any], overrides: Dict[str, Any]) -> Dict[str, Any]:
    tnved_list = _collect_tnved_list(all_data) or []
    tnved_list_over = overrides.get("g33_1_list")

    if not isinstance(tnved_list_over, (list, tuple)):
        tnved_list_over = overrides.get("g33_list")

    if isinstance(tnved_list_over, (list, tuple)):
        cleaned = []
        for x in tnved_list_over:
            s = "" if x in (None, "") else str(x).strip()
            if s:
                cleaned.append(s)
        tnved_list = cleaned

    return {"g33_1_list": tnved_list}

def compute_g34(all_data: Dict[str, Any], overrides: Dict[str, Any]) -> Dict[str, Any]:
    from graph import get_product_country, normalize_country, get_country_code
    tnved_list = _collect_tnved_list(all_data) or []
    tnved_list = sorted(tnved_list)  
    raw = get_product_country(all_data)
    g34_1_list_over = overrides.get("g34_1_list")
    if isinstance(g34_1_list_over, (list, tuple)):
        g34_1_list: List[str] = [str(x).strip() for x in g34_1_list_over]
    else:
        g34_1_list = []

        if isinstance(raw, dict) and tnved_list:
            for code in tnved_list:
                country_name = (
                    raw.get(code)
                    or raw.get(code.replace(" ", ""))
                    or ""
                )
                if country_name:
                    country_name = normalize_country(country_name)
                    country_code = get_country_code(country_name) or ""
                else:
                    country_code = ""
                g34_1_list.append(country_code)

        else:
            if isinstance(raw, set):
                raw = next(iter(sorted(raw))) if raw else ""
            if raw:
                raw = normalize_country(raw)
                base_code = get_country_code(raw) or ""
            else:
                base_code = ""

            if tnved_list:
                g34_1_list = [base_code] * len(tnved_list)
            else:
                g34_1_list = [base_code] if base_code else []

    return {
        "g34_1_list": g34_1_list
    }

def compute_g35(all_data: Dict[str, Any], overrides: Dict[str, Any]) -> Dict[str, Any]:
    from graph import get_brutto

    tnved_list = _collect_tnved_list(all_data) or []
    target_len = len(tnved_list) or 1

    try:
        brutto_map = get_brutto(all_data) or {}
    except Exception:
        brutto_map = {}

    over_list = overrides.get("g35_1_list")
    if not isinstance(over_list, (list, tuple)):
        over_list = overrides.get("g35_list")

    if isinstance(over_list, (list, tuple)):
        g35_list = [("" if v in (None, "") else str(v).strip()) for v in over_list]
    else:
        g35_list = []
        for code in tnved_list:
            val = brutto_map.get(code)
            g35_list.append("" if val in (None, "") else str(val))

    if len(g35_list) < target_len:
        g35_list += [""] * (target_len - len(g35_list))
    elif len(g35_list) > target_len:
        g35_list = g35_list[:target_len]

    return {"g35_1_list": g35_list}


def compute_g36(all_data: Dict[str, Any], overrides: Dict[str, Any]) -> Dict[str, Any]:
    default_g36_1 = "ОО"
    default_g36_2 = "ОО"
    default_g36_3 = "-"
    default_g36_4 = "ОО"
    return {
        "g36_1": overrides.get("g36_1", default_g36_1),
        "g36_2": overrides.get("g36_2", default_g36_2),
        "g36_3": overrides.get("g36_3", default_g36_3),
        "g36_4": overrides.get("g36_4", default_g36_4),
    }

def compute_g37(all_data: Dict[str, Any], overrides: Dict[str, Any]) -> Dict[str, Any]:
    tnved_list = _collect_tnved_list(all_data)
    g1_vals = compute_g1(all_data, overrides)
    base_g1_2 = (g1_vals.get("g1_2") or "").strip()
    g1_2 = (overrides.get("g1_2") or base_g1_2 or "").strip()

    over_list = overrides.get("g37_1_list")
    g37_1_list: List[str]
    base_proc = ""
    
    if isinstance(over_list, (list, tuple)):
        g37_1_list = [str(x).strip() for x in over_list]
        if tnved_list:
            need_len = len(tnved_list)
            if len(g37_1_list) < need_len:
                g37_1_list += [""] * (need_len - len(g37_1_list))
            elif len(g37_1_list) > need_len:
                g37_1_list = g37_1_list[:need_len]
    else:
        if g1_2:
            base_proc = f"{g1_2}"
        if tnved_list:
            g37_1_list = [base_proc] * len(tnved_list)
        else:
            g37_1_list = [base_proc]


    return {
        "g37_1_list": g37_1_list
    }

def compute_g38(all_data: Dict[str, Any], overrides: Dict[str, Any]) -> Dict[str, Any]:
    from graph import get_netto

    tnved_list = _collect_tnved_list(all_data) or []
    target_len = len(tnved_list) or 1

    try:
        netto_map = get_netto(all_data) or {}
    except Exception:
        netto_map = {}

    over_list = overrides.get("g38_1_list")
    if not isinstance(over_list, (list, tuple)):
        over_list = overrides.get("g38_list")

    if isinstance(over_list, (list, tuple)):
        g38_list = [("" if v in (None, "") else str(v).strip()) for v in over_list]
    else:
        g38_list = []
        for code in tnved_list:
            val = netto_map.get(code)
            g38_list.append("" if val in (None, "") else str(val))

    if len(g38_list) < target_len:
        g38_list += [""] * (target_len - len(g38_list))
    elif len(g38_list) > target_len:
        g38_list = g38_list[:target_len]

    return {"g38_1_list": g38_list}

def compute_g39(all_data: Dict[str, Any], overrides: Dict[str, Any]) -> Dict[str, Any]:
    tnved_list = _collect_tnved_list(all_data)
    target_len = len(tnved_list) or 1

    g39_1_list: List[str]
    over_list = overrides.get("g39_1_list")
    if isinstance(over_list, (list, tuple)):
        g39_1_list = [str(x).strip() for x in over_list]

        if len(g39_1_list) < target_len:
            g39_1_list += [""] * (target_len - len(g39_1_list))
        elif len(g39_1_list) > target_len:
            g39_1_list = g39_1_list[:target_len]
    else:
        g39_1_list = [""] * target_len

    return {
        "g39_1_list": g39_1_list,
    }

def compute_g40(all_data: Dict[str, Any], overrides: Dict[str, Any]) -> Dict[str, Any]:
    tnved_list = _collect_tnved_list(all_data)
    target_len = len(tnved_list) or 1  

    g40_1_list: List[str]
    over_list = overrides.get("g40_1_list")
    if isinstance(over_list, (list, tuple)):
        g40_1_list = [str(x).strip() for x in over_list]
        if len(g40_1_list) < target_len:
            g40_1_list += [""] * (target_len - len(g40_1_list))
        elif len(g40_1_list) > target_len:
            g40_1_list = g40_1_list[:target_len]
    else:
        g40_1_list = [""] * target_len

    return {
        "g40_1_list": g40_1_list,
    }

def compute_g41(all_data: Dict[str, Any], overrides: Dict[str, Any]) -> Dict[str, Any]:
    from graph import get_unit_tnved
    tnved_list = _collect_tnved_list(all_data) or []
    target_len = len(tnved_list)

    try:
        qty_by_tnved, unit_name_by_tnved, unit_code_by_tnved = get_unit_tnved(all_data)
    except Exception:
        qty_by_tnved, unit_name_by_tnved, unit_code_by_tnved = {}, {}, {}

    default_qty_list: List[str] = []
    default_unit_list: List[str] = []
    default_code_list: List[str] = []

    for code in tnved_list:
        qty_val = qty_by_tnved.get(code)
        qty_str = "" if qty_val in (None, "") else str(qty_val)
        default_qty_list.append(qty_str)

        unit_name = unit_name_by_tnved.get(code) or ""
        default_unit_list.append(unit_name)
        unit_code = unit_code_by_tnved.get(code) or ""

        if not unit_code:
            uname_up = unit_name.strip().upper()
            if qty_str not in ("", "0") and (uname_up == "ШТ" or uname_up == ""):
                unit_code = "796"

        default_code_list.append(unit_code)

    def apply_override(list_key: str, base_list: List[str]) -> List[str]:
        over = overrides.get(list_key)
        if isinstance(over, (list, tuple)):
            lst = [("" if v is None else str(v).strip()) for v in over]
            if len(lst) < target_len:
                lst += [""] * (target_len - len(lst))
            elif len(lst) > target_len:
                lst = lst[:target_len]
            return lst
        return base_list

    g41_1_list = apply_override("g41_1_list", default_qty_list)
    g41_2_list = apply_override("g41_2_list", default_unit_list)
    g41_3_list = apply_override("g41_3_list", default_code_list)

    return {
        "g41_1_list": g41_1_list,
        "g41_2_list": g41_2_list,
        "g41_3_list": g41_3_list,
    }

def compute_g42(all_data: Dict[str, Any], overrides: Dict[str, Any]) -> Dict[str, Any]:
    from graph import get_total_sum_tnved

    tnved_list = _collect_tnved_list(all_data) or []
    target_len = len(tnved_list)

    try:
        sums_map = get_total_sum_tnved(all_data) or {}
    except Exception:
        sums_map = {}

    default_prices: List[str] = []

    for code in tnved_list:
        val = sums_map.get(code)
        if isinstance(val, float):
            price_str = str(val.normalize())
        else:
            price_str = "" if val in (None, "") else str(val)
        default_prices.append(price_str)

    def apply_override(list_key: str, base_list: List[str]) -> List[str]:
        over = overrides.get(list_key)
        if isinstance(over, (list, tuple)):
            lst = [("" if v is None else str(v).strip()) for v in over]
            if len(lst) < target_len:
                lst += [""] * (target_len - len(lst))
            elif len(lst) > target_len:
                lst = lst[:target_len]
            return lst
        return base_list

    g42_1_list = apply_override("g42_1_list", default_prices)

    return {
        "g42_1_list": g42_1_list,
    }

def compute_g43(all_data: Dict[str, Any], overrides: Dict[str, Any]) -> Dict[str, Any]:
    tnved_list = _collect_tnved_list(all_data) or []
    target_len = len(tnved_list)
    base_val = overrides.get("g43_1", "1")
    base_val = ("" if base_val is None else str(base_val).strip()) or "1"

    default_list = [base_val] * target_len

    def apply_override(list_key: str, base_list):
        over = overrides.get(list_key)
        if isinstance(over, (list, tuple)):
            lst = [("" if v is None else str(v).strip()) for v in over]
            if len(lst) < target_len:
                lst += [""] * (target_len - len(lst))
            elif len(lst) > target_len:
                lst = lst[:target_len]
            return lst
        return base_list

    g43_1_list = apply_override("g43_1_list", default_list)

    return {
        "g43_1_list": g43_1_list,
    }

def compute_g44(all_data: Dict[str, Any], overrides: Dict[str, Any]) -> Dict[str, Any]:
    from graph import get_all_docx, get_transport, get_transport_type

    over_modes = overrides.get("g44_1_list")
    if isinstance(over_modes, (list, tuple)):
        return {
            "g44_1_list": list(over_modes),
            "g44_2_list": list(overrides.get("g44_2_list", [])),
            "g44_3_list": list(overrides.get("g44_3_list", [])),
            "g44_4_list": list(overrides.get("g44_4_list", [])),
            "g44_5_list": list(overrides.get("g44_5_list", [])),
        }
    if "g25_1" in overrides:
        g25_1_val = overrides["g25_1"]
    else:
        try:
            nums = get_transport(all_data)
            g25_1_val = get_transport_type(all_data, nums) or ""
        except Exception:
            g25_1_val = ""

    try:
        parts = get_all_docx(all_data, g25_1_val) or {}
    except Exception:
        parts = {}

    return {
        "g44_1_list": parts.get("mode_codes", []) or [],
        "g44_2_list": parts.get("kind_codes", []) or [],
        "g44_3_list": parts.get("names", []) or [],
        "g44_4_list": parts.get("numbers", []) or [],
        "g44_5_list": parts.get("dates_iso", []) or []
    }


def compute_g45(all_data: Dict[str, Any], overrides: Dict[str, Any]) -> Dict[str, Any]:
    from decimal import Decimal
    from graph import _to_decimal

    tnved_list = _collect_tnved_list(all_data) or []
    target_len = len(tnved_list) or 1

    def _norm_list(v, n: int) -> List[str]:
        if isinstance(v, (list, tuple)):
            lst = [("" if x in (None, "") else str(x).strip()) for x in v]
        else:
            lst = []
        if len(lst) < n:
            lst += [""] * (n - len(lst))
        elif len(lst) > n:
            lst = lst[:n]
        return lst

    over_list = overrides.get("g45_1_list")
    if isinstance(over_list, (list, tuple)):
        return {"g45_1_list": _norm_list(over_list, target_len)}

    if isinstance(overrides.get("g42_1_list"), (list, tuple)):
        g42_list = _norm_list(overrides.get("g42_1_list"), target_len)
    else:
        try:
            base42 = compute_g42(all_data, overrides)
            g42_list = _norm_list(base42.get("g42_1_list"), target_len)
        except Exception:
            g42_list = [""] * target_len

    rate_str = overrides.get("g23_1")
    if not rate_str:
        try:
            base23 = compute_g23(all_data, overrides)
            rate_str = base23.get("g23_1") or ""
        except Exception:
            rate_str = ""
    rate = _to_decimal(rate_str)
    if rate <= Decimal("0"):
        return {"g45_1_list": [""] * target_len}

    if overrides.get("g12_logistics") is not None or overrides.get("g12_insurance") is not None:
        add_total = _to_decimal(overrides.get("g12_logistics")) + _to_decimal(overrides.get("g12_insurance"))
    else:
        try:
            base12 = compute_g12(all_data, overrides)
            add_total = _to_decimal(base12.get("g12_logistics")) + _to_decimal(base12.get("g12_insurance"))
        except Exception:
            add_total = Decimal("0")
    add_total = add_total.quantize(Decimal("0.01")) if add_total != 0 else Decimal("0.00")

    try:
        base35 = compute_g35(all_data, overrides)
        g35_list = _norm_list(base35.get("g35_1_list"), target_len)
    except Exception:
        g35_list = [""] * target_len

    brutto_dec = [_to_decimal(x) for x in g35_list]
    total_brutto = sum([b for b in brutto_dec if b > 0], Decimal("0"))

    result: List[str] = []
    for i in range(target_len):
        inv_val = _to_decimal(g42_list[i])
        base_rub = (inv_val * rate).quantize(Decimal("0.01")) if inv_val > 0 else Decimal("0.00")

        share = Decimal("0.00")
        if add_total > 0 and total_brutto > 0 and brutto_dec[i] > 0:
            share = (add_total * brutto_dec[i] / total_brutto).quantize(Decimal("0.01"))

        item_total = (base_rub + share).quantize(Decimal("0.01"))
        result.append("" if item_total <= 0 else str(item_total))

    return {"g45_1_list": result}



def compute_g46(all_data: Dict[str, Any], overrides: Dict[str, Any]) -> Dict[str, Any]:
    from parser_cbrf import cb_rate
    from graph import get_any, _to_decimal

    tnved_list = _collect_tnved_list(all_data) or []
    target_len = len(tnved_list)
    over_list = overrides.get("g46_1_list")
    if isinstance(over_list, (list, tuple)):
        vals = [("" if v in (None, "") else str(v).strip()) for v in over_list]
        if len(vals) < target_len:
            vals += [""] * (target_len - len(vals))
        elif len(vals) > target_len:
            vals = vals[:target_len]
        return {"g46_1_list": vals}

    if "g45_1_list" in overrides:
        src45 = overrides.get("g45_1_list")
        if isinstance(src45, (list, tuple)):
            g45_list_raw = list(src45)
        else:
            g45_list_raw = []
    else:
        try:
            g45 = compute_g45(all_data, overrides)
            g45_list_raw = g45.get("g45_1_list") or []
        except Exception:
            g45_list_raw = []

    g45_list: List[str] = [("" if v in (None, "") else str(v).strip()) for v in g45_list_raw]
    if len(g45_list) < target_len:
        g45_list += [""] * (target_len - len(g45_list))
    elif len(g45_list) > target_len:
        g45_list = g45_list[:target_len]

    raw_date = get_any(all_data, [
        "declaration.Дата декларации",
        "declaration.date",
        "dt.Дата декларации",
        "dt.date"
    ])

    def _normalize_date_ddmmyyyy(d) -> str:
        if not d:
            return ""
        s = str(d).strip()
        m = re.fullmatch(r"(\d{4})-(\d{2})-(\d{2})", s)
        if m:
            y, mo, dd = m.groups()
            return f"{dd}.{mo}.{y}"
        m2 = re.fullmatch(r"(\d{2})\.(\d{2})\.(\d{4})", s)
        if m2:
            return s
        return s

    date_str = _normalize_date_ddmmyyyy(raw_date)
    if not date_str:
        return {"g46_1_list": [""] * target_len}

    try:
        usd_rate = cb_rate(date_str, "USD")
    except Exception:
        usd_rate = Decimal("0")

    if usd_rate <= Decimal("0"):
        return {"g46_1_list": [""] * target_len}

    result: List[str] = []
    for v in g45_list:
        base_val = _to_decimal(v)
        if base_val <= Decimal("0"):
            result.append("")
            continue
        try:
            stat_value = (base_val / usd_rate).quantize(Decimal("0.01"))
            result.append(str(stat_value))
        except Exception:
            result.append("")

    return {
        "g46_1_list": result,
    }

def compute_goods(all_data: Dict[str, Any], overrides: Dict[str, Any]) -> Dict[str, Any]:
    from graph import get_currency, _to_decimal, get_any
    from parser_cbrf import cb_rate
    from decimal import Decimal

    def _d(val) -> Decimal:
        try:
            return _to_decimal(val)
        except Exception:
            return Decimal("0")

    over_goods = overrides.get("goods_by_tnved")
    tnved_filter = str(overrides.get("goods_tnved_filter") or "").strip()

    if isinstance(over_goods, dict) and over_goods:
        goods_by_tnved: Dict[str, List[Dict[str, Any]]] = over_goods
        if tnved_filter:
            filtered: Dict[str, List[Dict[str, Any]]] = {}
            for code, items in goods_by_tnved.items():
                if str(code).startswith(tnved_filter):
                    filtered[str(code)] = items
            goods_by_tnved = filtered
        return {"goods_by_tnved": goods_by_tnved}

    invoice = (
        all_data.get("invoice")
        or all_data.get("invoice_json")
        or all_data.get("invoice_parsed")
        or {}
    )

    goods_src = None
    if isinstance(invoice, dict):
        goods_src = invoice.get("Товары") or invoice.get("goods") or invoice.get("items")

    if not isinstance(goods_src, list):
        return {"goods_by_tnved": {}}

    sender = invoice.get("Отправитель") if isinstance(invoice, dict) else None
    if not isinstance(sender, dict):
        sender = {}
    manufacturer_default = (sender.get("Название компании") or "").strip()

    goods_by_tnved: Dict[str, List[Dict[str, Any]]] = {}

    for idx, g in enumerate(goods_src):
        if not isinstance(g, dict):
            continue

        code = str(g.get("Код ТНВЭД") or "").strip()
        if not code:
            continue

        original_name = str(g.get("Наименование") or "").strip()
        tech_desc = str(g.get("Техническое описание") or "").strip()
        display_name = tech_desc or original_name

        trademark = str(g.get("Товарный знак") or "").strip()
        goods_mark = str(g.get("Марка") or "").strip() or "ОТСУТСТВУЕТ"
        goods_model = str(g.get("Модель") or "").strip() or "ОТСУТСТВУЕТ"
        goods_marking = str(g.get("Артикул") or "").strip() or "ОТСУТСТВУЕТ"

        qty = g.get("Количество") or ""
        currency = str(g.get("Валюта") or "").strip()
        if currency.lower() == "null":
            currency = ""

        price = _d(g.get("Цена"))
        qty_dec = _d(qty)
        invoiced_cost = _d(g.get("Стоимость"))
        if invoiced_cost <= 0:
            invoiced_cost = price * qty_dec

        item = {
            "index": idx,
            "tnved": code,
            "name": display_name,
            "original_name": original_name,
            "tech_desc": tech_desc,
            "manufacturer": manufacturer_default,
            "goods_trademark": trademark,
            "goods_mark": goods_mark,
            "goods_model": goods_model,
            "goods_marking": goods_marking,
            "qty": str(qty),
            "currency": currency,
            "invoiced_cost": str(invoiced_cost) if invoiced_cost != 0 else "",
        }

        goods_by_tnved.setdefault(code, []).append(item)

    if tnved_filter:
        filtered: Dict[str, List[Dict[str, Any]]] = {}
        for code, items in goods_by_tnved.items():
            if str(code).startswith(tnved_filter):
                filtered[code] = items
        goods_by_tnved = filtered

    return {"goods_by_tnved": goods_by_tnved}


def compute_graphs(all_data: Dict[str, Any],overrides: Optional[Dict[str, Any]]) -> Dict[str, Any]:

    overrides = overrides or {}
    graphs: Dict[str, Any] = {}
    graphs.update(compute_date_declararion(all_data, overrides))
    graphs.update(compute_g1(all_data, overrides))
    graphs.update(compute_g2(all_data, overrides))
    graphs.update(compute_g3(all_data, overrides))
    graphs.update(compute_g4(all_data, overrides))
    graphs.update(compute_g5(all_data, overrides))
    graphs.update(compute_g6(all_data, overrides))
    graphs.update(compute_g7(all_data, overrides))
    graphs.update(compute_g8(all_data, overrides))
    graphs.update(compute_g9(all_data, overrides))
    graphs.update(compute_g11(all_data, overrides))
    graphs.update(compute_g12(all_data, overrides))
    graphs.update(compute_g14(all_data, overrides))
    graphs.update(compute_g15(all_data, overrides))
    graphs.update(compute_g16(all_data, overrides))
    graphs.update(compute_g17(all_data, overrides))
    graphs.update(compute_g18(all_data, overrides)) 
    graphs.update(compute_g19(all_data, overrides)) 
    graphs.update(compute_g20(all_data, overrides)) 
    graphs.update(compute_g21(all_data, overrides)) 
    graphs.update(compute_g22(all_data, overrides)) 
    graphs.update(compute_g23(all_data, overrides)) 
    graphs.update(compute_g24(all_data, overrides)) 
    graphs.update(compute_g25(all_data, overrides)) 
    graphs.update(compute_g26(all_data, overrides)) 
    graphs.update(compute_g29(all_data, overrides)) 
    graphs.update(compute_g30(all_data, overrides)) 
    graphs.update(compute_g31(all_data, overrides)) 
    graphs.update(compute_g32(all_data, overrides)) 
    graphs.update(compute_g33(all_data, overrides)) 
    graphs.update(compute_g34(all_data, overrides)) 
    graphs.update(compute_g35(all_data, overrides)) 
    graphs.update(compute_g36(all_data, overrides)) 
    graphs.update(compute_g37(all_data, overrides)) 
    graphs.update(compute_g38(all_data, overrides)) 
    graphs.update(compute_g39(all_data, overrides)) 
    graphs.update(compute_g40(all_data, overrides)) 
    graphs.update(compute_g41(all_data, overrides)) 
    graphs.update(compute_g42(all_data, overrides)) 
    graphs.update(compute_g43(all_data, overrides)) 
    graphs.update(compute_g44(all_data, overrides)) 
    graphs.update(compute_g45(all_data, overrides)) 
    graphs.update(compute_g46(all_data, overrides)) 
    graphs.update(compute_goods(all_data, overrides))

    return graphs

######################## Вывод на React ########################
class GraphsOut(BaseModel):
    graphs: Dict[str, Any]

class GraphsUpdateIn(BaseModel):
    changes: Dict[str, Any]

@decl_router.get("/{decl_id}/graphs", response_model=GraphsOut)
def api_get_graphs(decl_id: int):
    all_data = build_all_data_for_decl(decl_id)
    overrides = get_overrides(decl_id)
    graphs = compute_graphs(all_data, overrides)
    graphs["document_id"] = f"declaration_{str(decl_id)}"
    return GraphsOut(graphs=graphs)

@decl_router.post("/{decl_id}/graphs", response_model=GraphsOut)
def api_update_graphs(decl_id: int, body: GraphsUpdateIn):
    all_data = build_all_data_for_decl(decl_id)
    overrides = get_overrides(decl_id) or {}
    changes = body.changes or {}
    for key, val in changes.items():
        if val in (None, "", [], {}):
            overrides.pop(key, None)
        else:
            overrides[key] = val

    auto_graphs = compute_graphs(all_data, overrides={})
    for key in list(overrides.keys()):
        try:
            if overrides.get(key) == auto_graphs.get(key):
                overrides.pop(key, None)
        except Exception:
            pass

    save_overrides(decl_id, overrides)
    graphs = compute_graphs(all_data, overrides)
    graphs["document_id"] = f"declaration_{decl_id}"
    return GraphsOut(graphs=graphs)

@graphs_router.get("/g30/by-tp")
def api_compute_g30_by_tp(tp_code: str):
    try:
      result = compute_g30(all_data={}, overrides={"g30_3": tp_code})
      return result
    except Exception as e:
      raise HTTPException(status_code=500, detail=str(e))
    
def _split_250(s: Any) -> List[str]:
    s = "" if s is None else str(s)
    if not s:
        return []
    return [s[i:i + 250] for i in range(0, len(s), 250)]

def _norm_list(payload: Dict[str, Any], list_key: str, scalar_key: str, n: int) -> List[str]:
    v = payload.get(list_key)
    if isinstance(v, (list, tuple)):
        lst = ["" if x is None else str(x) for x in v]
    else:
        sv = payload.get(scalar_key)
        lst = ["" if sv is None else str(sv)]

    if n <= 0:
        return lst

    if len(lst) < n:
        lst = lst + [""] * (n - len(lst))
    elif len(lst) > n:
        lst = lst[:n]
    return lst


def fill_ESADout_CU_with_gt(payload: Dict[str, Any]) -> ESADout_CU:
    g1_1 = payload.get("g1_1", "ИМ")
    g1_2 = payload.get("g1_2", "40")
    g1_3 = payload.get("g1_3", "ЭД")

    document_id_str = payload.get("document_id")
    document_id = ESADout_CU_DocumentID(document_id_str)

    declaration_dt = (
        payload.get("declaration_date")
        or (payload.get("declaration") or {}).get("Дата декларации")
        or ""
    )
    raw_tnved_list = payload.get("g33_1_list")
    if isinstance(raw_tnved_list, (list, tuple)):
        n = len(raw_tnved_list)
    else:
        n = 1

    g33_1_list = _norm_list(payload, "g33_1_list", "g33_1", n)
    g34_1_list = _norm_list(payload, "g34_1_list", "g34_1", n)
    g35_1_list = _norm_list(payload, "g35_1_list", "g35_1", n)
    g38_1_list = _norm_list(payload, "g38_1_list", "g38_1", n)

    g37_1_list = _norm_list(payload, "g37_1_list", "g37_1", n)
    g41_1_list = _norm_list(payload, "g41_1_list", "g41_1", n)
    g41_2_list = _norm_list(payload, "g41_2_list", "g41_2", n)
    g41_3_list = _norm_list(payload, "g41_3_list", "g41_3", n)

    g42_1_list = _norm_list(payload, "g42_1_list", "g42_1", n)
    g45_1_list = _norm_list(payload, "g45_1_list", "g45_1", n)
    g46_1_list = _norm_list(payload, "g46_1_list", "g46_1", n)

    g31_1_list = payload.get("g31_1_list")
    if isinstance(g31_1_list, (list, tuple)):
        g31_1_list = _norm_list(payload, "g31_1_list", "g31_1", n)
    else:
        one = payload.get("g31_1", "")
        g31_1_list = ["" if one is None else str(one)] * n

    raw_mode_codes = payload.get("g44_1_list")
    m = len(raw_mode_codes) if isinstance(raw_mode_codes, (list, tuple)) else 0

    g44_1_list = _norm_list(payload, "g44_1_list", "g44_1_list", m)
    g44_2_list = _norm_list(payload, "g44_2_list", "g44_2_list", m)
    g44_3_list = _norm_list(payload, "g44_3_list", "g44_3_list", m)
    g44_4_list = _norm_list(payload, "g44_4_list", "g44_4_list", m)
    g44_5_list = _norm_list(payload, "g44_5_list", "g44_5_list", m)
    g44_6_list = _norm_list(payload, "g44_6_list", "g44_6_list", m)
    g44_7_list = _norm_list(payload, "g44_7_list", "g44_7_list", m)
    g44_8_list = _norm_list(payload, "g44_8_list", "g44_8_list", m)

    presented_docs_common = []
    for j in range(m):
        mode = (g44_1_list[j] or "").strip()
        if not mode:
            continue

        kind = (g44_2_list[j] or "0").strip() or "0"

        doc_presenting = DocumentPresentingDetails(
            presented_document_mode_code=mode,
            doc_present_kind_code=kind,
        )

        presented_doc = ESADout_CUPresentedDocument(
            pr_document_name=g44_3_list[j],
            pr_document_number=g44_4_list[j],
            pr_document_date=g44_5_list[j],     
            presented_document_mode_code=mode,
            document_begin_actions_date=g44_6_list[j], 
            document_end_actions_date=g44_7_list[j],     
            record_id=g44_8_list[j],
            document_presenting_details=doc_presenting,
        )

        presented_docs_common.append(presented_doc)

    consignor = ESADout_CUConsignor(
        organization_name=payload.get("g2_3", ""),
        subject_address_details=SubjectAddressDetails(
            postal_code=payload.get("g2_6", ""),
            country_code=payload.get("g2_4", ""),
            country_name=payload.get("g2_5", ""),
            region=payload.get("g2_7", ""),
            city=payload.get("g2_8", ""),
            street_house=payload.get("g2_9", ""),
            house=payload.get("g2_10", ""),
        )
    )

    declarant = ESADout_CUDeclarant(
        organization_name=payload.get("g14_3", ""),
        rf_organization_features=RFOrganizationFeatures(
            ogrn=payload.get("g14_11", ""),
            inn=payload.get("g14_1", ""),
            kpp=payload.get("g14_2", ""),
        ),
        subject_address_details=SubjectAddressDetails(
            postal_code=payload.get("g14_6", ""),
            country_code=payload.get("g14_4", ""),
            country_name=payload.get("g14_5", ""),
            region=payload.get("g14_7", ""),
            city=payload.get("g14_8", ""),
            street_house=payload.get("g14_9", ""),
            house=payload.get("g14_10", ""),
        ),
    )

    border_office = BorderCustomsOffice(
        code=payload.get("g29_1", ""),  
        office_name=payload.get("g29_2",""),
        customs_country_code="643"
    )

    raw_ids = str(payload.get("g18_2") or payload.get("g21_2") or "").upper()
    country_code = str(payload.get("g18_3") or payload.get("g21_3") or "").upper()
    parts = [p.strip() for p in re.split(r"[;/,]+", raw_ids) if p.strip()]

    ru_transport_means = []

    for ident in parts:
        ru_transport_means.append(
            RUTransportMeans(
                transport_identifier=ident,
                transport_means_nationality_code=country_code,
            )
        )

    if not ru_transport_means and country_code:
        ru_transport_means.append(
            RUTransportMeans(
                transport_identifier="",
                transport_means_nationality_code=country_code,
            )
        )
    transport_means_qty = (
        payload.get("g18_1")
        or payload.get("g21_1")
        or (str(len(ru_transport_means)) if ru_transport_means else "")
    )

    departure_transport = ESADout_CUDepartureArrivalTransport(
        transport_mode_code=payload.get("g25_1", ""),   
        transport_means_quantity=transport_means_qty,
        ru_transport_means=ru_transport_means,
    )

    border_transport = ESADout_CUBorderTransport(
        transport_mode_code=payload.get("g25_1", ""),
        transport_means_quantity=payload.get("g21_1", ""),
    )

    consigment = ESADout_CUConsigment(
        container_indicator=payload.get("g19_1", ""),
        dispatch_country_code=payload.get("g15_1", ""),
        dispatch_country_name=payload.get("g15_2", ""),
        destination_country_code=payload.get("g17_1", ""),
        destination_country_name=payload.get("g17_2", ""),
        border_customs_office=border_office,
        departure_arrival_transport=departure_transport,
        border_transport=border_transport,
    )

    location_address = SubjectAddressDetails(
        country_code=payload.get("g30_country_code", ""),
        country_name=payload.get("g30_country_name", ""),
        region=payload.get("g30_region", ""),
        city=payload.get("g30_city", ""),
        street_house=payload.get("g30_street_house", ""),
        house="", 
    )

    goods_location = ESADout_CUGoodsLocation(
        information_type_code=payload.get("g30_1", ""),           
        customs_office=payload.get("g30_3", ""),                   
        customs_country_code=payload.get("g30_country_code", "RU"),
        location_name=payload.get("g30_svh_name", ""),                                         
        register_document_id_details=RegisterDocumentIdDetails(
            doc_id=payload.get("g30_license_number", "")
        ),
        address=location_address,
    )

    trade_country = payload.get("g11_1") or payload.get("g2_4") or ""
    delivery_terms = CUESADDeliveryTerms(
        delivery_place=payload.get("g20_2", ""),
        delivery_terms_string_code=payload.get("g20_1", ""),
    )
    contract_terms = ESADout_CUMainContractTerms(
        contract_currency_code=payload.get("g22_1", ""),
        contract_currency_rate=payload.get("g23_1", ""),
        total_invoice_amount=payload.get("g22_2", ""),
        trade_country_code=trade_country,
        deal_feature_code=payload.get("g24_2", ""),
        deal_nature_code=payload.get("g24_1", ""),
        cu_esad_delivery_terms=delivery_terms,
    )

    preferencii = Preferencii(
        customs_tax=payload.get("g36_1", ""),
        customs_duty=payload.get("g36_2", ""),
        excise=payload.get("g36_3", ""),
        rate=payload.get("g36_4", ""),
    )

    pallete_info_1 = PackagePalleteInformation(
        info_kind_code="", 
        pallete_code="PK",
        pallete_quantity="" 
    )

    goods_packaging = ESADGoodsPackaging(
        pakage_quantity=payload.get("g6_1", ""),
        pakage_type_code="1",
        package_pallete_information=[pallete_info_1],
    )

    goods_list: List[ESADout_CUGoods] = []
    goods_by_tnved = payload.get("goods_by_tnved") or {}
    if not isinstance(goods_by_tnved, dict):
        goods_by_tnved = {}

    for i in range(n):
        tnved_code = str(g33_1_list[i] or "")
        items_for_code = goods_by_tnved.get(tnved_code) or []

        group_desc_list = []

        if isinstance(items_for_code, list) and items_for_code:
            for j, good in enumerate(items_for_code):
                goods_group_qty = GoodsGroupQuantity(
                    goods_quantity=good.get("qty", ""),
                    measure_unit_qualifier_name=g41_2_list[i],
                    measure_unit_qualifier_code=g41_3_list[i],
                )

                goods_group_info = GoodsGroupInformation(
                    manufacturer=good.get("manufacturer", ""),
                    goods_mark=good.get("goods_mark", "ОТСУТСТВУЕТ"),
                    goods_model=good.get("goods_model", "ОТСУТСТВУЕТ"),
                    goods_marking=good.get("goods_marking", "ОТСУТСТВУЕТ"),
                    goods_group_quantity=goods_group_qty,
                    invoiced_cost=good.get("invoiced_cost", ""),
                )

                group_desc = GoodsGroupDescription(
                    goods_description=good.get("name", ""),
                    goods_group_information=goods_group_info,
                    group_num=str(j + 1),
                )
                group_desc_list.append(group_desc)
        else:
            goods_group_qty = GoodsGroupQuantity(
                goods_quantity="",
                measure_unit_qualifier_name=g41_2_list[i],
                measure_unit_qualifier_code=g41_3_list[i],
            )
            goods_group_info = GoodsGroupInformation(
                manufacturer="",
                goods_mark="ОТСУТСТВУЕТ",
                goods_model="ОТСУТСТВУЕТ",
                goods_marking="ОТСУТСТВУЕТ",
                serial_number="",
                goods_group_quantity=goods_group_qty,
                invoiced_cost="",
            )
            group_desc_list = [
                GoodsGroupDescription(
                    goods_description="",
                    goods_group_information=goods_group_info,
                    group_num="1",
                )
            ]

        suppl_quantity = SupplementaryGoodsQuantity(
            goods_quantity=g41_1_list[i],
            measure_unit_qualifier_name=g41_2_list[i],
            measure_unit_qualifier_code=g41_3_list[i],
        )

        customs_procedure_goods = ESADCustomsProcedure(
            main_customs_mode_code=g37_1_list[i],
            preceding_customs_mode_code="00",
            goods_transfer_feature="000",
        )

        goods = ESADout_CUGoods(
            goods_numeric=str(i + 1),
            goods_descriptions=_split_250(g31_1_list[i]),
            gross_weight_quantity=g35_1_list[i],
            net_weight_quantity=g38_1_list[i],
            invoiced_cost=g42_1_list[i], 
            customs_cost=g45_1_list[i],
            statistical_cost=g46_1_list[i],
            goods_tnved_code=tnved_code,
            intellect_property_sign="N",
            origin_country_code=g34_1_list[i],
            customs_cost_correct_method=payload.get("g43", ""),
            additional_sheet_count="1",
            goods_group_description=group_desc_list,  
            preferencii=preferencii,
            language_goods="RU",
            presented_documents=presented_docs_common,
            esad_goods_packaging=goods_packaging,
            esad_customs_procedure=customs_procedure_goods,
            supplementary_goods_quantity=suppl_quantity,
        )

        goods_list.append(goods)

    goods_shipment = ESADout_CUGoodsShipment(
        origin_country_name=payload.get("g16_2", ""),
        origin_country_code=payload.get("g16_1", ""),
        total_goods_number=str(n),
        total_package_number=str(payload.get("g6_1", "") or ""),
        total_sheet_number=str(n),
        total_cust_cost=str(payload.get("g12_1", "") or ""),
        cust_cost_currency_code=str(payload.get("g12_currency", "") or ""),
        consignor=consignor,
        consignee=ESADout_CUConsignee(equal_indicator="1"),
        financial_adjusting_responsible_person=ESADout_CUFinancialAdjustingResponsiblePerson(declarant_equal_flag="1"),
        declarant=declarant,
        consigment=consigment,
        main_contract_terms=contract_terms,
        goods_location=goods_location,
        goods_list=goods_list,
    )
    ece_doc_header_add_info = EECEDocHeaderAddInfo(
        e_doc_code="R.036",
        e_doc_date_time=declaration_dt,
        language_code="RU",
        source_country_code="RU",
        destination_country_code="RU",
    )

    esa_dout_cu = ESADout_CU(
        document_mode_id="1006107E",
        document_id=document_id,
        customs_procedure=CustomsProcedure(g1_1),
        customs_mode_code=CustomsModeCode(g1_2),
        electronic_document_sign=ElectronicDocumentSign(g1_3),
        recipient_country_code=RecipientCountryCode("RU"),
        goods_shipment=goods_shipment,
        ece_doc_header_add_info=ece_doc_header_add_info,
    )

    return esa_dout_cu


def _payload_from_graphs(graphs: Dict[str, Any]) -> Dict[str, Any]:
    payload: Dict[str, Any] = {}
    for k, v in graphs.items():
        payload[k] = v

    return payload


@decl_router.get("/{decl_id}/xml")
def api_get_declaration_xml(decl_id: int):
    try:
        all_data = build_all_data_for_decl(decl_id)
        overrides = get_overrides(decl_id) or {}
        graphs = compute_graphs(all_data, overrides)
        graphs["document_id"] = f"declaration_{str(decl_id)}"
        payload = _payload_from_graphs(graphs)
        esad = fill_ESADout_CU_with_gt(payload)
        xml_elem = esad.to_xml()
        xml_bytes = etree.tostring(
            xml_elem,
            encoding="utf-8",
            xml_declaration=True,
            pretty_print=True,
        )

        filename = f"declaration_{decl_id}.xml"
        try:
            user_id = get_declaration_user_id(decl_id)
        except Exception:
            user_id = None

        if user_id:
            try:
                file_id = add_file(
                    user_id=user_id,
                    filename=filename,
                    mime="application/xml",
                    file_bytes=xml_bytes,
                )
                update_declaration(decl_id, attached_file_id=file_id)
            except Exception as e:
                print(f"[XML] Ошибка сохранения файла для декларации {decl_id}: {e}")
        else:
            print(f"[XML] Не найден user_id для декларации {decl_id}, файл не привязан.")
        return Response(
            content=xml_bytes,
            media_type="application/xml",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    

app.include_router(auth_router)
app.include_router(users_router)
app.include_router(decl_router)
app.include_router(jobs_router)
app.include_router(admin_router)

app.include_router(files_router)
app.include_router(company_router)
app.include_router(tnved_router)
app.include_router(graphs_router)
app.include_router(debug_router)
app.include_router(misc_router)
